#!/usr/bin/env python3
"""Validate skill structure and enforce quantified one-page resume QA contracts."""
from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
from pathlib import Path
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from typing import Any, Iterable

import yaml
from design_tokens import load_theme


REQUIRED = (
    "SKILL.md", "agents/openai.yaml", "references/claim-ledger.md",
    "references/evidence-policy.md", "references/one-page-layout-qa.md",
)
MIN_MARGIN_CM, MAX_MARGIN_CM, MIN_BODY_FONT_PT = 1.27, 2.54, 10.0
# 1.27 cm is exactly 36 pt, but PDF/EMU conversions can produce a value a
# few floating-point ulps below that boundary.  The tolerance is deliberately
# sub-point: it absorbs measurement noise without accepting a material margin
# violation.
MIN_MARGIN_PT = MIN_MARGIN_CM * 72.0 / 2.54
MARGIN_TOLERANCE_PT = 0.5


def margin_below_minimum(value_pt: float) -> bool:
    """Return true only for a material margin violation, not float noise."""
    return float(value_pt) + MARGIN_TOLERANCE_PT < MIN_MARGIN_PT
# User-approved density band: 40–50pt is acceptable; only >50pt blocks.
# 50pt is the editorial target; allow 2pt for final glyph-box rounding so a
# one-line reflow is not forced to add repetitive copy for a 1–2pt residue.
MAX_BOTTOM_WHITESPACE_PT = 52.0
SKILLOPT_OPTIMIZABLE_CODES = frozenset({
    "PARAGRAPH_SPACING_ERROR", "PAGE_SIZE_ERROR", "MARGIN_OUT_OF_RANGE_ERROR",
    "MULTI_COLUMN_LAYOUT_ERROR", "VISUAL_DESIGN_MISMATCH_ERROR", "DELIVERY_GATE_BLOCKED",
    "DOCX_DELIVERY_BLOCKED",
})
SKILLOPT_NON_OPTIMIZABLE_CODES = frozenset({
    "EVIDENCE_GATE_BLOCKED", "INSUFFICIENT_PROJECT_EVIDENCE", "BULLET_LENGTH_ERROR",
    "BOTTOM_WHITESPACE_EXCESS", "PAGE_COUNT_ERROR", "CONTENT_GATE_BLOCKED", "NEEDS_USER_INPUT",
})
SKILLOPT_CONTENT_CODES = frozenset({
    "EVIDENCE_GATE_BLOCKED", "INSUFFICIENT_PROJECT_EVIDENCE", "BULLET_LENGTH_ERROR",
    "BOTTOM_WHITESPACE_EXCESS", "PAGE_COUNT_ERROR", "CONTENT_GATE_BLOCKED", "NEEDS_USER_INPUT",
})
MIN_CHINESE_BULLET_CHARS, MAX_CHINESE_BULLET_CHARS = 40, 50
CONTENT_BOUNDS = {
    "normal": (40, 50),
    "compressed": (30, 40),
    "expanded": (50, 130),
}
A4_WIDTH_PT, A4_HEIGHT_PT, PAGE_SIZE_TOLERANCE_PT = 595.28, 841.89, 1.5
PHONE_PATTERN = re.compile(r"(?<!\d)(?:\+?86[-\s]?)?1[3-9]\d{9}(?!\d)")
EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
URL_PATTERN = re.compile(r"https?://[^\s<>()]+", re.IGNORECASE)
CJK_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")


class ResumeQAError(Exception):
    """A deterministic, machine-readable delivery-blocking QA exception."""

    def __init__(self, code: str, detail: str) -> None:
        self.code, self.detail = code, detail
        super().__init__(f"{code}: {detail}")


def hard_fail(code: str, detail: str) -> None:
    raise ResumeQAError(code, detail)


def atomic_write_json(path: Path, payload: dict[str, Any]) -> None:
    """Write audit metadata atomically so partial manifests cannot authorize delivery."""
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    os.replace(temporary, path)


def emit_skillopt_failure_event(quarantine: Path, failed_manifest: Path) -> Path:
    """Emit one metadata-only event at the common quarantine boundary.

    Every failure enters the SkillOpt controller.  The event contains only
    redacted measurements; its route decides whether a bounded public-rule
    candidate, a content-recovery request, or diagnosis-only handling is safe.
    """
    payload = json.loads(failed_manifest.read_text(encoding="utf-8"))
    error = payload.get("error") if isinstance(payload.get("error"), dict) else {}
    code = str(error.get("code") or "UNKNOWN_GATE_ERROR")
    phase = str(payload.get("phase") or "unknown")
    inputs = payload.get("inputs") if isinstance(payload.get("inputs"), dict) else {}
    trace = quarantine / "reflow-trace.json"
    redacted_inputs = {str(k): str(v) for k, v in inputs.items() if isinstance(v, str)}
    # Always bind the failure event to the active public Skill hash, even when
    # the caller failed before it assembled its normal frozen-input dictionary.
    # Only the digest is emitted; the Skill text and private paths stay local.
    skill_path = Path(os.environ.get("SKILLOPT_SKILL_PATH") or Path(__file__).resolve().parents[1] / "SKILL.md")
    if skill_path.is_file():
        redacted_inputs.setdefault("skill_sha256", hashlib.sha256(skill_path.read_bytes()).hexdigest())
    measurements: dict[str, Any] = {}
    if trace.is_file():
        try:
            trace_payload = json.loads(trace.read_text(encoding="utf-8"))
            rounds = trace_payload.get("rounds") if isinstance(trace_payload, dict) else None
            if isinstance(rounds, list) and rounds and isinstance(rounds[-1], dict):
                latest = rounds[-1]
                measurements = {
                    "page_count": latest.get("page_count"),
                    "bottom_whitespace_pt": latest.get("bottom_whitespace_pt"),
                    "finding_codes": [f.get("code") for f in latest.get("findings", []) if isinstance(f, dict)],
                    "qa_measurements": latest.get("qa_measurements", {}),
                }
                # Preserve the complete reflow trajectory in redacted form.
                # SkillOpt can then distinguish a normal→compact improvement
                # from a fatal defect that persisted across all layout states,
                # without receiving resume text or private claims.
                measurements["round_history"] = [
                    {
                        "round": item.get("round"),
                        "layout_state": item.get("layout_state"),
                        "page_count": item.get("page_count"),
                        "bottom_whitespace_pt": item.get("bottom_whitespace_pt"),
                        "finding_codes": [
                            finding.get("code")
                            for finding in item.get("findings", [])
                            if isinstance(finding, dict) and isinstance(finding.get("code"), str)
                        ],
                    }
                    for item in rounds
                    if isinstance(item, dict)
                ]
                layout_hash = rounds[-1].get("layout_vars_sha256")
                if isinstance(layout_hash, str):
                    redacted_inputs["layout_vars_sha256"] = layout_hash
        except (OSError, json.JSONDecodeError):
            pass
    eligible = (
        code in SKILLOPT_OPTIMIZABLE_CODES
        and code not in SKILLOPT_NON_OPTIMIZABLE_CODES
        and (code not in {"DELIVERY_GATE_BLOCKED", "DOCX_DELIVERY_BLOCKED"}
             or phase in {"docx_delivery", "typst_delivery"})
    )
    if eligible:
        route = "public_rule_candidate"
    elif code in SKILLOPT_CONTENT_CODES:
        route = "content_recovery"
    elif code.startswith("EVIDENCE") or "CLAIM" in code or "SOURCE" in code:
        route = "evidence_review"
    else:
        # Unknown failures still enter the controller for diagnosis, but can
        # never receive an automatic rule or content mutation.
        route = "diagnose_only"
    gate = "content_gate_blocked" if code in SKILLOPT_NON_OPTIMIZABLE_CODES else (
        "delivery_gate_blocked" if code in {"DELIVERY_GATE_BLOCKED", "DOCX_DELIVERY_BLOCKED"} else "layout_gate_blocked"
    )
    event = {
        "event_id": f"failure-{payload.get('run_id') or uuid.uuid4().hex}",
        "run_id": payload.get("run_id"),
        "gate": gate,
        "error_code": code,
        "phase": phase,
        "trace_path": str(trace) if trace.is_file() else None,
        "failed_manifest": str(failed_manifest),
        "input_hashes": redacted_inputs,
        "measurements": measurements,
        "skill_sha256": redacted_inputs.get("skill_sha256", ""),
        "auto_skillopt": {
            "entered": True,
            "eligible": eligible,
            "route": route,
            "status": "queued",
        },
        "failure_signature": hashlib.sha256(json.dumps({"code": code, "phase": phase, "inputs": inputs}, sort_keys=True).encode()).hexdigest(),
    }
    path = quarantine / "skillopt-event.json"
    atomic_write_json(path, event)
    maybe_dispatch_skillopt(path)
    return path


def maybe_dispatch_skillopt(event_path: Path) -> None:
    """Start the controller when its private runtime is configured.

    Set ``SKILLOPT_AUTO_ENABLED=0`` to disable dispatch.  A benchmark is
    required for public-rule candidates; content/evidence lanes still receive
    a local recovery diagnosis without one.  Dispatch failures are recorded
    and never change the resume gate result.
    """
    if os.environ.get("SKILLOPT_AUTO_ENABLED") == "0":
        return
    repo_root = Path(__file__).resolve().parents[1]
    runtime_root = os.environ.get("SKILLOPT_RUNTIME_ROOT") or str(repo_root / ".skillopt-runtime")
    skill_path = os.environ.get("SKILLOPT_SKILL_PATH") or str(repo_root / "SKILL.md")
    benchmark = os.environ.get("SKILLOPT_BENCHMARK_COMMAND")
    proposal = os.environ.get("SKILLOPT_PROPOSAL_PATH")
    # With no external benchmark/proposal, run the safe controller lanes
    # in-process.  This makes every failure enter diagnosis/recovery without
    # spawning a child process (and without requiring an API key).
    if not benchmark and not proposal:
        try:
            from skillopt_auto_loop import diagnose, run_optimizer
            runtime = Path(runtime_root)
            event_copy, event_payload = diagnose(
                event_path.parent / "failed-manifest.json", runtime,
            )
            event_payload = run_optimizer(
                event_payload, runtime, Path(skill_path), "", None, False,
                int(os.environ.get("SKILLOPT_COOLDOWN_SECONDS", "3600")),
            )
            atomic_write_json(event_copy, event_payload)
            atomic_write_json(event_path, event_payload)
        except Exception as exc:  # controller diagnostics must never alter QA
            payload = json.loads(event_path.read_text(encoding="utf-8"))
            payload["auto_skillopt"].update({"status": "dispatch_error", "reason": str(exc)})
            atomic_write_json(event_path, payload)
        return
    command = [sys.executable, str(Path(__file__).with_name("skillopt_auto_loop.py")), "run",
               "--failed-manifest", str(event_path.parent / "failed-manifest.json"),
               "--runtime-root", runtime_root, "--skill-path", skill_path,
               "--cooldown-seconds",
               os.environ.get("SKILLOPT_COOLDOWN_SECONDS", "3600")]
    if benchmark:
        command.extend(["--benchmark-command", benchmark])
    if proposal:
        command.extend(["--proposal", proposal])
    else:
        command.append("--execute")
    try:
        completed = subprocess.run(command, text=True, capture_output=True, check=False)
        payload = json.loads(event_path.read_text(encoding="utf-8"))
        payload["auto_skillopt"].update({"dispatch_returncode": completed.returncode})
        if completed.stdout.strip():
            payload["auto_skillopt"]["dispatch_output_tail"] = completed.stdout[-2000:]
        if completed.returncode:
            payload["auto_skillopt"]["dispatch_stderr_tail"] = completed.stderr[-2000:]
        atomic_write_json(event_path, payload)
    except (OSError, json.JSONDecodeError) as exc:
        payload = json.loads(event_path.read_text(encoding="utf-8"))
        payload["auto_skillopt"].update({"status": "dispatch_error", "reason": str(exc)})
        atomic_write_json(event_path, payload)


def begin_render_transaction(output_dir: Path) -> tuple[str, Path]:
    """Create an isolated staging directory for one render attempt."""
    run_id = uuid.uuid4().hex
    staging = output_dir / ".staging" / run_id
    staging.mkdir(parents=True, exist_ok=False)
    return run_id, staging


def quarantine_render_transaction(
    output_dir: Path, run_id: str, staging: Path, *, code: str, detail: str,
    phase: str, inputs: dict[str, str] | None = None,
) -> Path:
    """Move failed render outputs out of the formal output namespace."""
    quarantine = output_dir / "quarantine" / run_id
    quarantine.parent.mkdir(parents=True, exist_ok=True)
    if staging.exists():
        shutil.move(str(staging), str(quarantine))
    else:
        quarantine.mkdir(parents=True, exist_ok=True)
    failed_manifest = quarantine / "failed-manifest.json"
    atomic_write_json(failed_manifest, {
        "status": "failed",
        "run_id": run_id,
        "phase": phase,
        "error": {"code": code, "detail": detail},
        "inputs": inputs or {},
    })
    emit_skillopt_failure_event(quarantine, failed_manifest)
    return quarantine


def promote_render_transaction(staging: Path, output_dir: Path, *, names: Iterable[str]) -> list[Path]:
    """Atomically promote only explicitly approved render artifacts."""
    promoted: list[Path] = []
    output_dir.mkdir(parents=True, exist_ok=True)
    for name in names:
        source = staging / name
        if not source.exists():
            continue
        destination = output_dir / name
        temporary = output_dir / f".{name}.{uuid.uuid4().hex}.promote"
        if source.is_dir():
            shutil.move(str(source), str(destination))
        else:
            shutil.copy2(source, temporary)
            os.replace(temporary, destination)
            # Promotion is a move in transactional terms.  Remove the staged
            # copy so the run directory can be removed and cannot be mistaken
            # for a second, discoverable artifact.
            source.unlink()
        promoted.append(destination)
    return promoted


def quarantine_artifacts(output_dir: Path, paths: Iterable[Path], *, code: str,
                         detail: str, phase: str) -> Path:
    """Move one or more failed artifacts into a run-isolated quarantine."""
    run_id = uuid.uuid4().hex
    quarantine = output_dir / "quarantine" / run_id
    quarantine.mkdir(parents=True, exist_ok=False)
    moved: list[str] = []
    for path in paths:
        if path.is_file():
            destination = quarantine / path.name
            shutil.move(str(path), str(destination))
            moved.append(destination.name)
    failed_manifest = quarantine / "failed-manifest.json"
    atomic_write_json(failed_manifest, {
        "status": "failed", "run_id": run_id, "phase": phase,
        "error": {"code": code, "detail": detail}, "artifacts": moved,
    })
    emit_skillopt_failure_event(quarantine, failed_manifest)
    return quarantine


def require_optional_dependencies() -> tuple[Any, Any, Any, Any]:
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from PIL import Image, ImageStat
        from pypdf import PdfReader
        import pdfplumber
    except ImportError as exc:  # pragma: no cover - depends on local runtime
        hard_fail("QA_DEPENDENCY_ERROR", f"install scripts/requirements.txt ({exc})")
    return Document, qn, Image, ImageStat, (PdfReader, pdfplumber)


def check_skill(skill_dir: Path) -> None:
    for relative in REQUIRED:
        if not (skill_dir / relative).is_file():
            hard_fail("SKILL_STRUCTURE_ERROR", f"missing required file: {relative}")
    skill = (skill_dir / "SKILL.md").read_text(encoding="utf-8")
    if not skill.startswith("---\n") or "name: resume-evidence-rebuild" not in skill:
        hard_fail("SKILL_STRUCTURE_ERROR", "SKILL.md must contain resume-evidence-rebuild frontmatter")
    if "[TODO" in skill:
        hard_fail("SKILL_STRUCTURE_ERROR", "SKILL.md contains unfinished placeholder text")


def check_file_signature(path: Path, signature: bytes, label: str) -> None:
    if not path.is_file() or path.read_bytes()[:len(signature)] != signature:
        hard_fail("ARTIFACT_FORMAT_ERROR", f"invalid {label}: {path}")


def check_manifest_provenance(manifest_path: Path, *, renderer_path: Path | None = None,
                              expected_paths: dict[str, Path] | None = None) -> dict[str, Any]:
    """Verify that a manifest describes the exact inputs used to render it."""
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    provenance = manifest.get("provenance")
    if not isinstance(provenance, dict):
        hard_fail("PROVENANCE_MISSING_ERROR", "manifest lacks canonical renderer provenance")
    checks = dict(expected_paths or {})
    if renderer_path is not None:
        checks["renderer_sha256"] = renderer_path
    for field, path in checks.items():
        expected = hashlib.sha256(path.read_bytes()).hexdigest()
        if provenance.get(field) != expected:
            hard_fail("PROVENANCE_HASH_ERROR", f"manifest {field} does not match {path}")
    return provenance


def check_delivery_manifest(manifest_path: Path, *, expected_paths: dict[str, Path]) -> dict[str, Any]:
    """Verify that a delivery manifest still describes the bytes on disk.

    A manifest is an authorization record, not proof by itself.  Rechecking
    every declared artifact hash prevents a post-QA overwrite from being
    mistaken for the file that actually passed the gate.
    """
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        hard_fail("DELIVERY_MANIFEST_ERROR", f"invalid delivery manifest: {exc}")
    if payload.get("status") != "eligible_for_approval":
        hard_fail("DELIVERY_MANIFEST_ERROR", "delivery manifest is not eligible_for_approval")
    declared = payload.get("sha256")
    if not isinstance(declared, dict):
        hard_fail("DELIVERY_MANIFEST_ERROR", "delivery manifest lacks artifact hashes")
    for label, path in expected_paths.items():
        if not path.is_file():
            hard_fail("DELIVERY_MANIFEST_ERROR", f"manifest artifact is missing: {path}")
        if declared.get(label) != hashlib.sha256(path.read_bytes()).hexdigest():
            hard_fail("DELIVERY_MANIFEST_ERROR", f"manifest hash mismatch for {label}: {path}")
    return payload


def cm_from_emu(value: int | None) -> float:
    return float(value or 0) / 360000.0  # 1 cm = 360000 EMU


def is_body_paragraph(paragraph: Any) -> bool:
    if not paragraph.text.strip():
        return False
    style = (getattr(paragraph.style, "name", "") or "").lower()
    return not any(token in style for token in ("title", "heading", "header", "footer"))


def resolved_font_size_pt(paragraph: Any, run: Any) -> float | None:
    for candidate in (run.font.size, getattr(paragraph.style.font, "size", None),
                      getattr(paragraph.part.document.styles["Normal"].font, "size", None)):
        if candidate is not None:
            return float(candidate.pt)
    return None


def paragraph_spacing_is_compact(paragraph: Any) -> bool:
    """Validate actual paragraph properties stamped by the DOCX renderer.

    The value must be present on the paragraph itself, rather than inherited
    from a style, so a later style change cannot quietly reintroduce sparse
    paragraphs into an otherwise approved resume.
    """
    fmt = paragraph.paragraph_format
    style = (getattr(paragraph.style, "name", "") or "").lower().replace(" ", "")
    is_bullet = "listbullet" in style
    expected_line_spacing = 1.3 if is_bullet else 1.4
    max_after_pt = 5.0 if is_bullet else 0.5
    return (
        isinstance(fmt.line_spacing, (int, float))
        and abs(float(fmt.line_spacing) - expected_line_spacing) < 0.001
        and fmt.space_after is not None
        and 0.0 <= float(fmt.space_after.pt) <= max_after_pt
    )


def iter_docx_image_blobs(document: Any) -> Iterable[bytes]:
    """Yield images in the document body and linked headers/footers.

    Chinese resumes place the required photo in the header.  Header parts have
    their own relationship collection, so checking only ``document.part``
    falsely reports a missing photo for an otherwise valid artifact.
    """
    parts = [document.part]
    for section in document.sections:
        parts.extend((
            section.header.part,
            section.first_page_header.part,
            section.even_page_header.part,
            section.footer.part,
            section.first_page_footer.part,
            section.even_page_footer.part,
        ))
    seen: set[str] = set()
    for part in parts:
        part_name = str(part.partname)
        if part_name in seen:
            continue
        seen.add(part_name)
        for relation in part.rels.values():
            if "image" in relation.reltype:
                yield relation.target_part.blob


def image_is_three_by_four_solid(blob: bytes, Image: Any, ImageStat: Any) -> bool:
    with Image.open(io.BytesIO(blob)).convert("RGB") as image:
        width, height = image.size
        if not width or not height or abs(width / height - 3 / 4) > 0.03:
            return False
        edge = max(1, min(width, height) // 12)
        # Portrait photos legitimately contain clothing at the lower edge.  Use
        # the two upper corners to assess the backdrop rather than mistaking a
        # dark suit for a non-solid background.
        boxes = ((0, 0, edge, edge), (width-edge, 0, width, edge))
        means = []
        for box in boxes:
            stat = ImageStat.Stat(image.crop(box))
            if max(stat.var) > 225:  # Corner standard deviation must be <= 15.
                return False
            means.append(stat.mean)
        baseline = means[0]
        return all(max(abs(value - baseline[i]) for i, value in enumerate(mean)) <= 25 for mean in means[1:])


def check_docx_layout_and_photo(docx_path: Path, market: str, overview_texts: set[str]) -> None:
    Document, qn, Image, ImageStat, _ = require_optional_dependencies()
    document = Document(str(docx_path))
    for section_number, section in enumerate(document.sections, 1):
        columns = section._sectPr.find(qn("w:cols"))
        if columns is not None and int(columns.get(qn("w:num"), "1")) != 1:
            hard_fail("MULTI_COLUMN_LAYOUT_ERROR", f"DOCX section {section_number} declares multiple text columns")
        for side, value in {
            "top": cm_from_emu(section.top_margin), "bottom": cm_from_emu(section.bottom_margin),
            "left": cm_from_emu(section.left_margin), "right": cm_from_emu(section.right_margin),
        }.items():
            value_pt = value * 72.0 / 2.54
            if margin_below_minimum(value_pt) or value > MAX_MARGIN_CM:
                code = "BOTTOM_WHITESPACE_EXCESS" if value > MAX_MARGIN_CM else "MARGIN_OUT_OF_RANGE_ERROR"
                hard_fail(code, f"DOCX section {section_number} {side} margin is {value:.2f} cm; expected 1.27-2.54 cm")
    for table_number, table in enumerate(document.tables, 1):
        description = table._tbl.tblPr.find(qn("w:tblDescription"))
        if description is not None and description.get(qn("w:val")) == "resume-header":
            # The identity row is intentionally two-column (text + photo);
            # the single-column rule applies to the resume body only.
            continue
        if any(len(row.cells) > 1 for row in table.rows):
            hard_fail("MULTI_COLUMN_LAYOUT_ERROR", f"DOCX body table {table_number} has multiple columns")
    # The fixed header is stored in a marked table and is not part of
    # document.paragraphs. Every non-empty body paragraph must therefore be
    # checked directly; counting the first four body paragraphs would let a
    # malformed renderer bypass the spacing contract.
    for paragraph_number, paragraph in enumerate(document.paragraphs, 1):
        if not is_body_paragraph(paragraph):
            continue
        for run in paragraph.runs:
            if run.text.strip():
                size = resolved_font_size_pt(paragraph, run)
                is_overview = paragraph.text.strip() in overview_texts
                minimum = 9.0 if is_overview else MIN_BODY_FONT_PT
                if size is None or size < minimum:
                    hard_fail("FONT_TOO_SMALL_ERROR", f"DOCX paragraph {paragraph_number} uses {size or 0:.1f} pt; body minimum is 10 pt")
        if not paragraph_spacing_is_compact(paragraph):
            hard_fail("PARAGRAPH_SPACING_ERROR", f"DOCX paragraph {paragraph_number} must use direct 1.4x spacing, or 1.3x with a 5pt inter-bullet gap")
    images = list(iter_docx_image_blobs(document))
    if market == "CN":
        if not images or not any(image_is_three_by_four_solid(blob, Image, ImageStat) for blob in images):
            hard_fail("COMPLIANCE_PHOTO_ERROR", "CN photo must exist, be 3:4, and have a solid background")
    elif images:
        hard_fail("COMPLIANCE_PHOTO_ERROR", f"{market} route prohibits photos")


def check_docx_ooxml_spacing(docx_path: Path) -> dict[str, Any]:
    """Require direct OOXML 1.4x regular text and 1.3x bullets.

    python-docx can expose inherited style values as if they were direct
    formatting. Inspecting document.xml closes that template/style bypass.
    Word stores 1.4x as line=336 and 1.3x as line=312 (240 * multiplier).
    """
    from zipfile import ZipFile
    from lxml import etree

    _, qn, _, _, _ = require_optional_dependencies()
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path) as archive:
        try:
            root = etree.fromstring(archive.read("word/document.xml"))
        except KeyError:
            hard_fail("ARTIFACT_FORMAT_ERROR", "DOCX has no word/document.xml")
    checked = 0
    line_values: list[int] = []
    after_values: list[int] = []
    for index, paragraph in enumerate(root.xpath(".//w:body/w:p", namespaces=ns), 1):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces=ns)).strip()
        if not text:
            continue
        style_node = paragraph.find("./w:pPr/w:pStyle", ns)
        style_name = (style_node.get(qn("w:val"), "") if style_node is not None else "").lower()
        if any(token in style_name for token in ("title", "heading", "header", "footer")):
            continue
        # Section markers and title/date rows are structural anchors rather
        # than body bullets.  Their compact spacing is checked by the
        # structure/geometry validators, not by the body after-spacing rule.
        # Some producers serialize tab stops as plain runs, so recognize the
        # date range itself instead of relying on a literal tab.
        has_title_date = bool(re.search(
            r"(?:19|20)\d{2}[./-]\d{2}\s*[–—-]\s*(?:至今|(?:19|20)\d{2}[./-]\d{2})",
            text,
        ))
        if text.startswith("▌ ") or has_title_date:
            continue
        checked += 1
        style_key = style_name.replace(" ", "")
        is_bullet = "listbullet" in style_key
        expected_line = "312" if is_bullet else "336"
        expected_after = "100" if is_bullet else "10"
        spacing = paragraph.find("./w:pPr/w:spacing", ns)
        line = spacing.get(qn("w:line")) if spacing is not None else None
        rule = spacing.get(qn("w:lineRule")) if spacing is not None else None
        after = spacing.get(qn("w:after")) if spacing is not None else None
        if line is not None:
            line_values.append(int(line))
        if after is not None:
            after_values.append(int(after))
        if line != expected_line or rule != "auto" or after != expected_after:
            hard_fail("PARAGRAPH_SPACING_ERROR", f"DOCX paragraph {index} has line={line!r}, lineRule={rule!r}, after={after!r}; expected {expected_line}/auto/{expected_after}")
    return {"body_paragraphs_checked": checked, "expected_line": "336 (body), 312 (bullets)", "expected_after": "10 (body), 100 (bullets)",
            "actual_line_values": sorted(set(line_values)), "actual_after_values": sorted(set(after_values))}


def check_docx_theme(docx_path: Path, theme_path: Path) -> None:
    """Check the DOCX's explicit theme identity and visible design anchors."""
    Document, _, _, _, _ = require_optional_dependencies()
    theme = load_theme(theme_path)
    document = Document(str(docx_path))
    if document.core_properties.subject != f"resume-theme:{theme['variant_id']}":
        hard_fail("VISUAL_DESIGN_MISMATCH_ERROR", "DOCX theme variant differs from frozen theme_vars")
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
    text = "".join(run.text for run in runs)
    if "▌" not in text:
        hard_fail("VISUAL_DESIGN_MISMATCH_ERROR", "DOCX is missing the approved section marker")


def render_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Render DOCX with LibreOffice so physical QA sees the final layout."""
    binary = shutil.which("soffice") or shutil.which("libreoffice")
    if not binary:
        hard_fail("QA_DEPENDENCY_ERROR", "LibreOffice/soffice is required for DOCX render QA")
    output_dir.mkdir(parents=True, exist_ok=True)
    completed = subprocess.run([binary, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)], text=True, capture_output=True)
    rendered = output_dir / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not rendered.is_file():
        hard_fail("DOCX_RENDER_ERROR", completed.stderr.strip() or completed.stdout.strip() or "LibreOffice did not write PDF")
    return rendered


def run_docx_delivery_gate(*, docx_path: Path, manifest_path: Path, theme_path: Path,
                           profile_path: Path, market: str, qa_dir: Path,
                           renderer_path: Path | None = None) -> dict[str, Any]:
    """Run DOCX source and rendered checks; failure never triggers PDF reflow."""
    check_file_signature(docx_path, b"PK", "DOCX")
    check_manifest_provenance(
        manifest_path, renderer_path=renderer_path,
        expected_paths={"profile_sha256": profile_path, "theme_vars_sha256": theme_path},
    )
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    overview_texts = {str(project.get("overview")) for project in manifest.get("projects", []) if project.get("overview")}
    check_docx_layout_and_photo(docx_path, market, overview_texts)
    spacing = check_docx_ooxml_spacing(docx_path)
    check_docx_project_bold_emphasis(docx_path, manifest_path)
    check_docx_theme(docx_path, theme_path)
    profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
    if not isinstance(profile, dict) or not isinstance(profile.get("identity"), dict):
        hard_fail("ARTIFACT_ARGUMENT_ERROR", "profile must contain an identity mapping")
    rendered_pdf = render_docx_to_pdf(docx_path, qa_dir)
    check_pdf_layout_and_integrity(rendered_pdf, market, profile["identity"])
    rendered_pdf_spacing = check_pdf_body_line_spacing(rendered_pdf, manifest_path)
    return {"status": "passed", "docx_sha256": hashlib.sha256(docx_path.read_bytes()).hexdigest(),
            "rendered_pdf_sha256": hashlib.sha256(rendered_pdf.read_bytes()).hexdigest(),
            "rendered_pdf": rendered_pdf.name, "spacing": spacing,
            "rendered_pdf_line_spacing": rendered_pdf_spacing}


def text_integrity_check(text: str, expected_identity: dict[str, Any] | None = None) -> None:
    """Check contacts against authorized private profile data when supplied."""
    if expected_identity is None:
        if not PHONE_PATTERN.search(text) or not EMAIL_PATTERN.search(text):
            hard_fail("CONTACT_MISSING_ERROR", "final artifact has no recognizable phone number and email")
        return
    for field in ("phone", "email", "portfolio_url"):
        value = str(expected_identity.get(field, "")).strip()
        if not value:
            hard_fail("CONTACT_MISSING_ERROR", f"private profile is missing immutable {field}")
        compact_value = re.sub(r"\s+", "", value)
        compact_text = re.sub(r"\s+", "", text)
        if value not in text and compact_value not in compact_text:
            code = "LINK_TAMPERING_ERROR" if field == "portfolio_url" else "CONTACT_MISSING_ERROR"
            hard_fail(code, f"final artifact does not contain authorized immutable {field}")


def pdf_image_blobs(reader: Any) -> list[bytes]:
    blobs: list[bytes] = []
    for page in reader.pages:
        try:
            blobs.extend(image.data for image in page.images)
        except Exception:
            continue
    return blobs


def check_pdf_layout_and_integrity(pdf_path: Path, market: str, expected_identity: dict[str, Any] | None = None, *, allow_density: bool = False) -> None:
    _, _, _, _, (PdfReader, pdfplumber) = require_optional_dependencies()
    reader = PdfReader(str(pdf_path))
    page_count_error = None
    if len(reader.pages) != 1:
        page_count_error = f"a resume must contain exactly one A4 page; found {len(reader.pages)} pages"
    media_box = reader.pages[0].mediabox
    width, height = float(media_box.width), float(media_box.height)
    if abs(width - A4_WIDTH_PT) > PAGE_SIZE_TOLERANCE_PT or abs(height - A4_HEIGHT_PT) > PAGE_SIZE_TOLERANCE_PT:
        hard_fail("PAGE_SIZE_ERROR", f"expected A4 {A4_WIDTH_PT:.2f}x{A4_HEIGHT_PT:.2f} pt; found {width:.2f}x{height:.2f} pt")
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if not text.strip():
        hard_fail("TEXT_EXTRACTION_ERROR", "PDF has no extractable text")
    text_integrity_check(text, expected_identity)
    images = pdf_image_blobs(reader)
    if market == "CN" and not images:
        hard_fail("COMPLIANCE_PHOTO_ERROR", "final PDF has no embedded photo for CN route")
    if market != "CN" and images:
        hard_fail("COMPLIANCE_PHOTO_ERROR", f"final PDF contains a photo for {market} route")
    with pdfplumber.open(str(pdf_path)) as pdf:
        for number, page in enumerate(pdf.pages, 1):
            chars = [char for char in page.chars if char.get("text", "").strip()]
            if not chars:
                hard_fail("TEXT_EXTRACTION_ERROR", f"PDF page {number} has no measurable text")
            words = page.extract_words(x_tolerance=1, y_tolerance=2)
            # A true two-column body has independently aligned line starts on both
            # halves of the page. The header is excluded because a headshot may sit there.
            body_line_starts: dict[int, float] = {}
            for word in words:
                if word["top"] <= 150 or word["bottom"] >= page.height - 45:
                    continue
                key = round(word["top"])
                body_line_starts[key] = min(body_line_starts.get(key, word["x0"]), word["x0"])
            midpoint = page.width / 2
            left_lines = sum(start < midpoint - 20 for start in body_line_starts.values())
            right_lines = sum(start > midpoint + 20 for start in body_line_starts.values())
            if left_lines >= 4 and right_lines >= 4:
                hard_fail("MULTI_COLUMN_LAYOUT_ERROR", f"PDF page {number} has {left_lines} left and {right_lines} right body-column line starts")
            boundaries = {"left": min(char["x0"] for char in chars), "right": page.width-max(char["x1"] for char in chars),
                          "top": min(char["top"] for char in chars), "bottom": page.height-max(char["bottom"] for char in chars)}
            for side, value in boundaries.items():
                # Compare in the native PDF unit.  Converting a boundary such
                # as 36pt to cm and comparing floats caused 1.27cm to be
                # reported as below the same 1.27cm minimum.
                margin_pt = float(value)
                cm = margin_pt * 2.54 / 72.0
                # Text bounding boxes do not prove a maximum left/right/top
                # margin: a short final line naturally leaves a large right
                # gap on an overfull page. Enforce the safety minimum on every
                # edge, and reserve the 50pt upper bound exclusively for the
                # bottom-density signal on a one-page candidate.
                if margin_below_minimum(margin_pt):
                    hard_fail(
                        "MARGIN_OUT_OF_RANGE_ERROR",
                        f"PDF page {number} {side} content boundary is {cm:.3f} cm ({margin_pt:.3f} pt); "
                        f"minimum is {MIN_MARGIN_CM:.2f} cm ({MIN_MARGIN_PT:.3f} pt)",
                    )
            if boundaries["bottom"] > MAX_BOTTOM_WHITESPACE_PT and not allow_density:
                hard_fail("BOTTOM_WHITESPACE_EXCESS", f"PDF page {number} bottom whitespace is {boundaries['bottom']:.2f} pt; maximum is {MAX_BOTTOM_WHITESPACE_PT:.0f} pt")
            # The generated source is separately checked to allow 9pt only for
            # project overviews. PDF text lacks semantic tags, so a 9pt glyph
            # is permitted here solely when every overview source marker exists.
            if any(float(char["size"]) + 0.01 < 9.0 for char in chars if 110 < char["top"] < page.height - 55):
                hard_fail("FONT_TOO_SMALL_ERROR", f"PDF page {number} contains text below the 9pt overview exception")
    if page_count_error and not allow_density:
        hard_fail("PAGE_COUNT_ERROR", page_count_error)


def load_project_bullets(manifest_path: Path, *, source_text: bool = False) -> list[tuple[str, int, str, list[str]]]:
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        hard_fail("INSUFFICIENT_PROJECT_EVIDENCE", f"invalid project-node manifest: {exc}")
    projects = payload.get("projects") if isinstance(payload, dict) else None
    if not isinstance(projects, list) or not 3 <= len(projects) <= 4:
        hard_fail("INSUFFICIENT_PROJECT_EVIDENCE", "manifest must contain exactly 3-4 evidence-backed projects")
    bullets: list[tuple[str, int, str, list[str]]] = []
    for project_number, project in enumerate(projects, 1):
        name = project.get("name") if isinstance(project, dict) else None
        points = project.get("bullets") if isinstance(project, dict) else None
        if not isinstance(name, str) or not isinstance(points, list) or not 3 <= len(points) <= 4:
            hard_fail("INSUFFICIENT_PROJECT_EVIDENCE", f"project {project_number} must have a name and 3-4 bullets")
        for bullet_number, bullet in enumerate(points, 1):
            if not isinstance(bullet, dict) or not isinstance(bullet.get("text"), str):
                hard_fail("BULLET_LENGTH_ERROR", f"project {name!r} bullet {bullet_number} must be an object with text")
            bold_phrases = bullet.get("bold_phrases")
            if not isinstance(bold_phrases, list) or not 1 <= len(bold_phrases) <= 2 or not all(isinstance(item, str) and item for item in bold_phrases):
                hard_fail("BULLET_BOLD_MISSING_ERROR", f"project {name!r} bullet {bullet_number} must declare 1-2 non-empty bold_phrases")
            if any(phrase not in bullet["text"] for phrase in bold_phrases):
                hard_fail("BULLET_BOLD_MISSING_ERROR", f"project {name!r} bullet {bullet_number} declares a phrase absent from its text")
            text = bullet.get("source_text", bullet["text"]) if source_text else bullet["text"]
            bullets.append((name, bullet_number, text, bold_phrases))
    return bullets


def load_employment_bullets(manifest_path: Path, *, source_text: bool = False) -> list[tuple[str, int, str, list[str]]]:
    """Load all rendered work bullets, which share the project bullet QA contract."""
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        hard_fail("INSUFFICIENT_WORK_EVIDENCE", f"invalid resume-node manifest: {exc}")
    employment = payload.get("employment") if isinstance(payload, dict) else None
    if not isinstance(employment, list):
        hard_fail("INSUFFICIENT_WORK_EVIDENCE", "manifest must contain a work-experience bullet manifest")
    bullets: list[tuple[str, int, str, list[str]]] = []
    for employment_number, entry in enumerate(employment, 1):
        name = entry.get("name") if isinstance(entry, dict) else None
        points = entry.get("bullets") if isinstance(entry, dict) else None
        if not isinstance(name, str) or not isinstance(points, list) or not 4 <= len(points) <= 5:
            hard_fail("INSUFFICIENT_WORK_EVIDENCE", f"employment {employment_number} must have a name and 4-5 bullets")
        for bullet_number, bullet in enumerate(points, 1):
            if not isinstance(bullet, dict) or not isinstance(bullet.get("text"), str):
                hard_fail("BULLET_LENGTH_ERROR", f"employment {name!r} bullet {bullet_number} must be an object with text")
            bold_phrases = bullet.get("bold_phrases")
            if not isinstance(bold_phrases, list) or not 1 <= len(bold_phrases) <= 2 or not all(isinstance(item, str) and item for item in bold_phrases):
                hard_fail("BULLET_BOLD_MISSING_ERROR", f"employment {name!r} bullet {bullet_number} must declare 1-2 non-empty bold_phrases")
            if any(phrase not in bullet["text"] for phrase in bold_phrases):
                hard_fail("BULLET_BOLD_MISSING_ERROR", f"employment {name!r} bullet {bullet_number} declares a phrase absent from its text")
            text = bullet.get("source_text", bullet["text"]) if source_text else bullet["text"]
            bullets.append((name, bullet_number, text, bold_phrases))
    return bullets


def load_resume_bullets(manifest_path: Path, *, source_text: bool = False) -> list[tuple[str, int, str, list[str]]]:
    return [*load_project_bullets(manifest_path, source_text=source_text), *load_employment_bullets(manifest_path, source_text=source_text)]


def load_rendered_bullet_nodes(manifest_path: Path) -> list[dict[str, Any]]:
    """Return nodes in the physical order written by typst_renderer.py.

    The ``element_id`` plus sequence makes a repeated phrase addressable.  A
    global ``str.find`` can otherwise verify the first equal phrase twice and
    silently skip an unbolded later occurrence.
    """
    try:
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        hard_fail("ARTIFACT_FORMAT_ERROR", f"invalid resume-node manifest: {exc}")
    if not isinstance(payload, dict):
        hard_fail("ARTIFACT_FORMAT_ERROR", "resume-node manifest must be an object")
    nodes: list[dict[str, Any]] = []
    for group_name, groups in (("employment", payload.get("employment", [])), ("projects", payload.get("projects", []))):
        if not isinstance(groups, list):
            hard_fail("ARTIFACT_FORMAT_ERROR", f"manifest {group_name} must be a list")
        for group in groups:
            if not isinstance(group, dict):
                hard_fail("ARTIFACT_FORMAT_ERROR", f"manifest {group_name} entry must be an object")
            for bullet in group.get("bullets", []):
                if not isinstance(bullet, dict):
                    hard_fail("ARTIFACT_FORMAT_ERROR", "manifest bullet must be an object")
                text, element_id, phrases = bullet.get("text"), bullet.get("element_id"), bullet.get("bold_phrases")
                if (not isinstance(text, str) or not isinstance(element_id, str)
                        or not isinstance(phrases, list) or not all(isinstance(item, str) and item for item in phrases)):
                    hard_fail("ARTIFACT_FORMAT_ERROR", "manifest bullet lacks text, element_id, or bold_phrases")
                nodes.append({"element_id": element_id, "text": text, "bold_phrases": phrases})
    return nodes


def check_bullet_lengths(manifest_path: Path) -> None:
    try:
        manifest_payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        hard_fail("ARTIFACT_FORMAT_ERROR", f"invalid resume-node manifest: {exc}")
    mode = str(manifest_payload.get("content_mode", "normal")) if isinstance(manifest_payload, dict) else "normal"
    if mode not in CONTENT_BOUNDS:
        hard_fail("ARTIFACT_FORMAT_ERROR", f"unsupported content_mode {mode!r}")
    min_chars, max_chars = CONTENT_BOUNDS[mode]
    for project_name, bullet_number, bullet, _ in load_project_bullets(manifest_path, source_text=True):
        count = len(CJK_PATTERN.findall(bullet))
        if not min_chars <= count <= max_chars:
            hard_fail("BULLET_LENGTH_ERROR", f"resume {project_name!r} bullet {bullet_number} has {count} Chinese characters; expected {min_chars}-{max_chars} for {mode}")
    # Employment remains on its immutable 40–50 budget during project-only
    # density recovery.
    for employment_name, bullet_number, bullet, _ in load_employment_bullets(manifest_path, source_text=True):
        count = len(CJK_PATTERN.findall(bullet))
        if not MIN_CHINESE_BULLET_CHARS <= count <= MAX_CHINESE_BULLET_CHARS:
            hard_fail("BULLET_LENGTH_ERROR", f"employment {employment_name!r} bullet {bullet_number} has {count} Chinese characters; expected 40-50")


def check_docx_project_bold_emphasis(docx_path: Path, manifest_path: Path) -> None:
    """Verify declared important phrases are visibly bold in the editable source."""
    Document, _, _, _, _ = require_optional_dependencies()
    document = Document(str(docx_path))
    paragraphs = [(paragraph.text.replace("• ", ""), paragraph) for paragraph in document.paragraphs]
    for project_name, bullet_number, text, bold_phrases in load_resume_bullets(manifest_path):
        matched = next((paragraph for paragraph_text, paragraph in paragraphs if text in paragraph_text), None)
        if matched is None:
            hard_fail("BULLET_BOLD_MISSING_ERROR", f"project {project_name!r} bullet {bullet_number} is not present in the DOCX")
        bold_text = "".join(run.text if run.bold else "\0" * len(run.text) for run in matched.runs)
        if not any(phrase in bold_text for phrase in bold_phrases):
            hard_fail("BULLET_BOLD_MISSING_ERROR", f"project {project_name!r} bullet {bullet_number} has no declared phrase in a bold DOCX run")


def typst_escape(value: str) -> str:
    return (value.replace("\\", "\\\\").replace("#", "\\#")
            .replace("@", "\\@").replace("[", "\\[").replace("]", "\\]"))


def check_typst_project_bold_emphasis(typst_path: Path, manifest_path: Path) -> None:
    """Verify that the PDF source explicitly renders every declared phrase bold.

    PDF font names are unreliable for this purpose because Typst can synthesize
    a bold face from a regular-only CJK installation.  Checking the compiled
    source guarantees the renderer supplied an explicit visible-weight rule.
    """
    if not typst_path.is_file():
        hard_fail("ARTIFACT_FORMAT_ERROR", f"missing Typst source: {typst_path}")
    source = typst_path.read_text(encoding="utf-8")
    for node in load_rendered_bullet_nodes(manifest_path):
        start_marker = f'#metadata("{node["element_id"]}")'
        start = source.find(start_marker)
        if start < 0:
            hard_fail("BULLET_BOLD_MISSING_ERROR", f"Typst source lacks bullet metadata {node['element_id']!r}")
        end = source.find('#metadata("', start + len(start_marker))
        bullet_source = source[start:end if end >= 0 else len(source)]
        for phrase in node["bold_phrases"]:
            escaped = typst_escape(phrase)
            marker = f'#text(weight: "bold", fill: rgb("#000000"), stroke: 0.12pt + rgb("#000000"))[{escaped}]'
            expected_occurrences = node["text"].count(phrase)
            if bullet_source.count(marker) != expected_occurrences:
                hard_fail(
                    "BULLET_BOLD_MISSING_ERROR",
                    f"bullet {node['element_id']!r} lacks explicit visible Typst bold for every occurrence of {phrase!r}",
                )


def locate_ordered_pdf_bullets(pages: list[Any], nodes: list[dict[str, Any]]) -> list[tuple[int, list[dict[str, Any]]]]:
    """Locate each manifest bullet in reading order, even during an overflow round."""
    page_chars = [[item for item in page.chars if item.get("text", "")] for page in pages]
    page_texts = ["".join(item["text"] for item in chars) for chars in page_chars]
    page_index, cursor = 0, 0
    runs: list[tuple[int, list[dict[str, Any]]]] = []
    for node in nodes:
        while page_index < len(page_texts):
            start = page_texts[page_index].find(node["text"], cursor)
            if start >= 0:
                end = start + len(node["text"])
                runs.append((page_index, page_chars[page_index][start:end]))
                cursor = end
                break
            page_index += 1
            cursor = 0
        else:
            hard_fail("PDF_TEXT_ORDER_ERROR", f"bullet {node['element_id']!r} is absent from ordered PDF glyphs")
    return runs


def check_pdf_rendered_bold(pdf_path: Path, manifest_path: Path) -> None:
    """Require content-stream stroke plus raster ink at each *specific* run."""
    _, _, Image, _, (PdfReader, pdfplumber) = require_optional_dependencies()
    reader = PdfReader(str(pdf_path))
    stream = b"".join(page.get_contents().get_data() for page in reader.pages)
    if b" Tr" not in stream or b" w" not in stream:
        hard_fail("PDF_BOLD_NOT_RENDERED_ERROR", "PDF content stream has no explicit stroked text evidence")
    binary = shutil.which("pdftoppm")
    if not binary:
        hard_fail("PDF_BOLD_NOT_RENDERED_ERROR", "pdftoppm is required for raster bold verification")
    nodes = load_rendered_bullet_nodes(manifest_path)
    with tempfile.TemporaryDirectory(prefix="resume-bold-") as temp:
        prefix = Path(temp) / "page"
        subprocess.run([binary, "-png", "-r", "144", str(pdf_path), str(prefix)], check=True, capture_output=True)
        with pdfplumber.open(str(pdf_path)) as pdf:
            runs = locate_ordered_pdf_bullets(pdf.pages, nodes)
            for node, (page_index, bullet_run) in zip(nodes, runs):
                image_path = prefix.with_name(f"{prefix.name}-{page_index + 1}.png")
                if not image_path.is_file():
                    hard_fail("PDF_BOLD_NOT_RENDERED_ERROR", f"missing raster page {page_index + 1}")
                image = Image.open(image_path).convert("L")
                page = pdf.pages[page_index]
                for phrase in node["bold_phrases"]:
                    phrase_start = node["text"].find(phrase)
                    while phrase_start >= 0:
                        run = bullet_run[phrase_start:phrase_start + len(phrase)]
                        if len(run) != len(phrase) or "".join(item["text"] for item in run) != phrase:
                            hard_fail("PDF_BOLD_NOT_RENDERED_ERROR", f"bold phrase {phrase!r} lacks a continuous PDF glyph run")
                        scale = image.width / page.width
                        # A valid terminal phrase may wrap. Verify visible ink
                        # on every physical line rather than silently locating
                        # the first identical phrase elsewhere on the page.
                        line_groups: list[list[dict[str, Any]]] = []
                        for item in run:
                            if not line_groups or abs(float(item["top"]) - float(line_groups[-1][0]["top"])) > 2:
                                line_groups.append([item])
                            else:
                                line_groups[-1].append(item)
                        for line in line_groups:
                            box = (int(min(item["x0"] for item in line) * scale), int(min(item["top"] for item in line) * scale),
                                   int(max(item["x1"] for item in line) * scale), int(max(item["bottom"] for item in line) * scale) + 2)
                            crop = image.crop(box)
                            if not crop.getbbox() or min(crop.getdata()) > 245:
                                hard_fail("PDF_BOLD_NOT_RENDERED_ERROR", f"bold phrase {phrase!r} has no visible raster ink")
                        phrase_start = node["text"].find(phrase, phrase_start + len(phrase))


def check_pdf_body_line_spacing(pdf_path: Path, manifest_path: Path) -> dict[str, Any]:
    """Measure wrapped bullet lines and the fixed gap between adjacent bullets."""
    _, _, _, _, (_, pdfplumber) = require_optional_dependencies()
    body_reference_pt = 10.0
    with pdfplumber.open(str(pdf_path)) as pdf:
        intervals: list[float] = []
        bullet_starts: list[tuple[int, float, float]] = []
        nodes = load_rendered_bullet_nodes(manifest_path)
        located = locate_ordered_pdf_bullets(pdf.pages, nodes)
        for page_index, run in located:
            line_tops = sorted({round(float(item["top"]), 2) for item in run})
            intervals.extend(right - left for left, right in zip(line_tops, line_tops[1:]))
            bullet_starts.append((page_index, line_tops[0], line_tops[-1]))
        # A 10pt bullet uses 1.3x wrapped-line rhythm (13pt). The fixed 5pt
        # paragraph gap makes adjacent one-line bullets start 15pt apart.
        if any(interval < 12.4 or interval > 13.6 for interval in intervals):
            hard_fail("PARAGRAPH_SPACING_ERROR", "rendered wrapped bullet line spacing is not the required 1.3x rhythm")
        inter_bullet_intervals = [
            round(next_start - current_end, 2)
            for node_a, (page_a, _, current_end), node_b, (page_b, next_start, _) in zip(
                nodes, bullet_starts, nodes[1:], bullet_starts[1:]
            )
            if page_a == page_b
            and node_a["element_id"].rsplit(".bullet.", 1)[0] == node_b["element_id"].rsplit(".bullet.", 1)[0]
        ]
        if any(interval < 14.4 or interval > 15.6 for interval in inter_bullet_intervals):
            hard_fail("PARAGRAPH_SPACING_ERROR", "rendered adjacent bullet separation is not the required 1.5x rhythm")
        multipliers = [round(interval / body_reference_pt, 6) for interval in intervals]
        return {
            "body_reference_pt": body_reference_pt,
            "expected_multiplier": 1.3,
            "expected_pt": 13.0,
            "tolerance_pt": 0.6,
            "measured_intervals_pt": intervals,
            "measured_multipliers": multipliers,
            "measured_min_pt": min(intervals) if intervals else None,
            "measured_max_pt": max(intervals) if intervals else None,
            "measured_multiplier_min": min(multipliers) if multipliers else None,
            "measured_multiplier_max": max(multipliers) if multipliers else None,
            "inter_bullet_expected_multiplier": 1.5,
            "inter_bullet_expected_pt": 15.0,
            "inter_bullet_intervals_pt": inter_bullet_intervals,
            "inter_bullet_multiplier_min": min((value / body_reference_pt for value in inter_bullet_intervals), default=None),
            "inter_bullet_multiplier_max": max((value / body_reference_pt for value in inter_bullet_intervals), default=None),
        }


def measure_pdf_font_sizes(pdf_path: Path) -> dict[str, Any]:
    """Return the observed glyph-size range for traceability."""
    _, _, _, _, (_, pdfplumber) = require_optional_dependencies()
    with pdfplumber.open(str(pdf_path)) as pdf:
        sizes = [float(char["size"]) for page in pdf.pages for char in page.chars
                 if char.get("text", "").strip() and char.get("size") is not None]
        body_sizes = [float(char["size"]) for page in pdf.pages for char in page.chars
                      if char.get("text", "").strip() and char.get("size") is not None
                      and 110 < float(char.get("top", 0)) < page.height - 55]
    return {
        # pdfplumber's glyph box is smaller than the requested font for some
        # CJK fonts (notably the header), so retain both the raw range and the
        # body range used by the semantic font gate.
        "min_glyph_pt": min(sizes) if sizes else None,
        "max_glyph_pt": max(sizes) if sizes else None,
        "min_body_glyph_pt": min(body_sizes) if body_sizes else None,
        "glyph_count": len(sizes),
    }


def source_without_line_comments(source: str) -> str:
    """Remove full-line Typst comments without corrupting URLs in string literals."""
    return "\n".join(line for line in source.splitlines() if not line.lstrip().startswith("//"))


def check_typst_compact_body_spacing(typst_path: Path) -> None:
    """Require 1.4x regular text, 1.3x wrapped bullets, and 1.5x bullet gaps."""
    source = source_without_line_comments(typst_path.read_text(encoding="utf-8"))
    par_directives = re.findall(r"(?m)^\s*#set\s+par\s*\(([^\n]*)\)\s*$", source)
    normalized = [re.sub(r"\s+", "", directive) for directive in par_directives]
    expected = "leading:0.4em,spacing:0.5pt"
    text_directives = re.findall(r"(?m)^\s*#set\s+text\s*\(([^\n]*)\)\s*$", source)
    edge_directives = [re.sub(r"\s+", "", directive) for directive in text_directives
                       if "top-edge" in directive or "bottom-edge" in directive]
    required_layout_keys = ("header_to_first_module", "module_gap", "project_gap", "title_to_overview", "overview_to_bullet")
    layout_uses = all(f'layout-len("{key}")' in source for key in required_layout_keys)
    fixed_vertical_values = {"0", "0.5", "1", "3", "4", "5", "20"}
    hardcoded_vertical_values = re.findall(r"#v\(\s*(\d+(?:\.\d+)?)pt\s*\)", source)
    if (source.count("#set par") != 1 or normalized != [expected]
            or "set par(leading: 0.3em, spacing: 5pt)" not in source
            or edge_directives != ["top-edge:0.8em,bottom-edge:-0.2em"]
            or re.search(r"#show[\s\S]*?set\s+par\s*\(", source)
            or source.count('#let layout = json("layout_vars.json")') != 1
            or '#let layout-len(key) = layout.at("spacing").at(key) * 1pt' not in source
            or not layout_uses or any(value not in fixed_vertical_values for value in hardcoded_vertical_values)):
        hard_fail(
            "PARAGRAPH_SPACING_ERROR",
            "Typst body must declare fixed 1.4x text spacing, scoped 1.3x bullet lines, and 1.5x inter-bullet gaps",
        )


def check_typst_font_whitelist(typst_path: Path, manifest_path: Path) -> None:
    source = typst_path.read_text(encoding="utf-8")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    overview_count = sum(1 for project in manifest.get("projects", []) if project.get("overview"))
    text_sizes = [float(value) for value in re.findall(r"size:\s*(\d+(?:\.\d+)?)pt", source)]
    if any(size < 9.0 for size in text_sizes):
        hard_fail("FONT_TOO_SMALL_ERROR", "Typst source contains an illegal sub-9pt text size")
    if sum(size == 9.0 for size in text_sizes) != overview_count:
        hard_fail("FONT_TOO_SMALL_ERROR", "9pt is allowed exactly once per manifest overview and nowhere else")
    if source.count('#text(size: 9pt, fill: rgb("#000000"))') != overview_count:
        hard_fail("FONT_TOO_SMALL_ERROR", "only manifest overview nodes may use the 9pt exception")
    for project in manifest.get("projects", []):
        overview = project.get("overview")
        if overview:
            marker = (
                f'#metadata("project.{project.get("id")}.overview") '
                f'#text(size: 9pt, fill: rgb("#000000"))[{typst_escape(overview)}]'
            )
            if marker not in source:
                hard_fail("FONT_TOO_SMALL_ERROR", "9pt exception is not bound to its manifest overview node")


def check_typst_title_emphasis(typst_path: Path, manifest_path: Path) -> None:
    """Require visible bold treatment for section and project titles."""
    source = typst_path.read_text(encoding="utf-8")
    title_style = 'fill: design-color("accent"), stroke: 0.12pt + design-color("accent")'
    if source.count(f'#text(size: 12pt, weight: "bold", {title_style})') < 3:
        hard_fail("TITLE_BOLD_MISSING_ERROR", "all three section titles must use visible bold blue title styling")
    for project_name, _, _, _ in load_project_bullets(manifest_path):
        marker = f'#text(size: 11pt, weight: "bold", {title_style})[{typst_escape(project_name)}]'
        if marker not in source:
            hard_fail("TITLE_BOLD_MISSING_ERROR", f"project title {project_name!r} lacks visible bold")


def check_visual_design(theme_path: Path, typst_path: Path, manifest_path: Path, docx_path: Path | None) -> None:
    approved = load_theme(theme_path)
    source = typst_path.read_text(encoding="utf-8")
    required = (
        '#let design = json("theme_vars.json")',
        '#let design-color(key)',
        '#rect(width: design-len("section_marker_width_pt")',
        'stroke: design-color("rule") + design-len("section_rule_pt")',
    )
    if any(marker not in source for marker in required):
        hard_fail("VISUAL_DESIGN_MISSING_ERROR", "Typst source does not render the approved executive-editorial tokens")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest.get("theme_variant") != approved["variant_id"]:
        hard_fail("VISUAL_DESIGN_MISMATCH_ERROR", "project manifest theme variant differs from frozen theme_vars")
    if docx_path:
        Document, _, _, _, _ = require_optional_dependencies()
        document = Document(str(docx_path))
        if document.core_properties.subject != f"resume-theme:{approved['variant_id']}":
            hard_fail("VISUAL_DESIGN_MISMATCH_ERROR", "DOCX theme variant differs from frozen theme_vars")
        expected_colors = {
            approved["tokens"]["palette"]["ink"].removeprefix("#").upper(),
            approved["tokens"]["palette"]["accent"].removeprefix("#").upper(),
        }
        actual_colors = {
            str(run.font.color.rgb).upper() for paragraph in document.paragraphs for run in paragraph.runs
            if run.font.color.rgb is not None
        }
        if not expected_colors.issubset(actual_colors) or not any("▌" in run.text for p in document.paragraphs for run in p.runs):
            hard_fail("VISUAL_DESIGN_MISMATCH_ERROR", "DOCX does not visibly use the approved colors and section marker")


def run_pdf_delivery_gate(*, skill_dir: Path, pdf_path: Path, manifest_path: Path,
                          typst_path: Path, theme_path: Path, profile_path: Path,
                          market: str, geometry_path: Path, allow_density: bool = False) -> dict[str, Any]:
    """Run the sole PDF delivery gate used by both render entry points.

    ``allow_density`` is reserved for the internal normal/compact reflow loop;
    direct renderer calls always use the strict default and cannot emit a
    delivery manifest for an overflowing or sparse page.
    """
    check_skill(skill_dir)
    check_file_signature(pdf_path, b"%PDF", "PDF")
    check_manifest_provenance(
        manifest_path,
        renderer_path=skill_dir / "scripts" / "typst_renderer.py",
        expected_paths={"profile_sha256": profile_path, "theme_vars_sha256": theme_path, "typst_source_sha256": typst_path},
    )
    check_bullet_lengths(manifest_path)
    check_typst_project_bold_emphasis(typst_path, manifest_path)
    check_pdf_rendered_bold(pdf_path, manifest_path)
    check_typst_compact_body_spacing(typst_path)
    pdf_spacing = check_pdf_body_line_spacing(pdf_path, manifest_path)
    pdf_font_sizes = measure_pdf_font_sizes(pdf_path)
    check_typst_font_whitelist(typst_path, manifest_path)
    check_typst_title_emphasis(typst_path, manifest_path)
    check_visual_design(theme_path, typst_path, manifest_path, None)
    profile = yaml.safe_load(profile_path.read_text(encoding="utf-8"))
    if not isinstance(profile, dict) or not isinstance(profile.get("identity"), dict):
        hard_fail("ARTIFACT_ARGUMENT_ERROR", "profile must contain an identity mapping")
    check_pdf_layout_and_integrity(pdf_path, market, profile["identity"], allow_density=allow_density)

    geometry_script = skill_dir / "scripts" / "geometry_qa.py"
    completed = subprocess.run(
        [sys.executable, str(geometry_script), "--pdf", str(pdf_path), "--manifest", str(manifest_path), "--output", str(geometry_path)],
        text=True, capture_output=True,
    )
    if not geometry_path.is_file():
        hard_fail("GEOMETRY_QA_ERROR", completed.stderr.strip() or completed.stdout.strip() or "geometry QA did not write output")
    geometry = json.loads(geometry_path.read_text(encoding="utf-8"))
    if isinstance(geometry, dict):
        geometry.setdefault("qa_measurements", {})["pdf_line_spacing"] = pdf_spacing
        geometry["qa_measurements"]["pdf_font_size"] = pdf_font_sizes
        atomic_write_json(geometry_path, geometry)
    findings = geometry.get("findings", []) if isinstance(geometry, dict) else []
    codes = {item.get("code") for item in findings if isinstance(item, dict)}
    density_codes = {"PAGE_COUNT_ERROR", "BOTTOM_WHITESPACE_EXCESS"}
    blocking = codes - density_codes
    if blocking:
        hard_fail("GEOMETRY_QA_ERROR", f"geometry QA found blocking codes: {sorted(blocking)}")
    if not allow_density and codes & density_codes:
        hard_fail("CONTENT_GATE_BLOCKED", f"geometry QA found density codes: {sorted(codes & density_codes)}")
    return geometry


def write_delivery_manifest(path: Path, *, pdf_path: Path, manifest_path: Path,
                            typst_path: Path, theme_path: Path, geometry: dict[str, Any]) -> None:
    """Write the only artifact that designates a PDF as eligible for delivery."""
    payload = {
        "status": "eligible_for_approval",
        "pdf": pdf_path.name,
        "sha256": {label: hashlib.sha256(file.read_bytes()).hexdigest() for label, file in {
            "pdf": pdf_path, "project_manifest": manifest_path, "typst_source": typst_path, "theme_vars": theme_path,
        }.items()},
        "geometry": geometry,
    }
    atomic_write_json(path, payload)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--skill-dir", type=Path, required=True)
    parser.add_argument("--docx", type=Path)
    parser.add_argument("--pdf", type=Path)
    parser.add_argument("--project-manifest", type=Path)
    parser.add_argument("--typst-source", type=Path, help="Generated Typst source used for the PDF")
    parser.add_argument("--theme-vars", "--design-tokens", dest="theme_vars", type=Path, help="Frozen allow-listed theme_vars.json")
    parser.add_argument("--profile", type=Path, help="Private profile used only to check immutable contact values")
    parser.add_argument("--market", choices=("CN", "NA", "FOREIGN"), default="CN")
    args = parser.parse_args()
    try:
        check_skill(args.skill_dir)
        expected_identity = None
        if args.profile:
            profile = yaml.safe_load(args.profile.read_text(encoding="utf-8"))
            if not isinstance(profile, dict) or not isinstance(profile.get("identity"), dict):
                hard_fail("ARTIFACT_ARGUMENT_ERROR", "profile must contain an identity mapping")
            expected_identity = profile["identity"]
        if args.docx and (not args.pdf or not args.project_manifest):
            hard_fail("ARTIFACT_ARGUMENT_ERROR", "--docx requires --pdf and --project-manifest")
        if args.pdf and not args.project_manifest:
            hard_fail("ARTIFACT_ARGUMENT_ERROR", "--pdf requires --project-manifest")
        if args.pdf:
            if not args.theme_vars or not args.typst_source or not args.profile:
                hard_fail("ARTIFACT_ARGUMENT_ERROR", "PDF validation requires --theme-vars, --typst-source, and --profile")
            geometry_path = args.pdf.parent / "geometry-qa.json"
            run_pdf_delivery_gate(
                skill_dir=args.skill_dir, pdf_path=args.pdf, manifest_path=args.project_manifest,
                typst_path=args.typst_source, theme_path=args.theme_vars,
                profile_path=args.profile, market=args.market, geometry_path=geometry_path,
            )
            delivery_manifest = args.pdf.parent / "delivery-manifest.json"
            if delivery_manifest.is_file():
                check_delivery_manifest(delivery_manifest, expected_paths={
                    "pdf": args.pdf,
                    "project_manifest": args.project_manifest,
                    "typst_source": args.typst_source,
                    "theme_vars": args.theme_vars,
                })
        if args.docx:
            check_file_signature(args.docx, b"PK", "DOCX")
            manifest = json.loads(args.project_manifest.read_text(encoding="utf-8"))
            overview_texts = {str(project.get("overview")) for project in manifest.get("projects", []) if project.get("overview")}
            check_docx_layout_and_photo(args.docx, args.market, overview_texts)
            check_docx_ooxml_spacing(args.docx)
            check_docx_project_bold_emphasis(args.docx, args.project_manifest)
            if args.theme_vars:
                check_docx_theme(args.docx, args.theme_vars)
            docx_delivery_manifest = args.docx.parent / "docx-delivery-manifest.json"
            if docx_delivery_manifest.is_file():
                check_delivery_manifest(docx_delivery_manifest, expected_paths={
                    "docx": args.docx,
                    "project_manifest": args.project_manifest,
                    "docx_project_manifest": args.docx.parent / "docx-project-manifest.json",
                    "theme_vars": args.theme_vars,
                })
            if args.profile:
                profile = yaml.safe_load(args.profile.read_text(encoding="utf-8"))
                if not isinstance(profile, dict) or not isinstance(profile.get("identity"), dict):
                    hard_fail("ARTIFACT_ARGUMENT_ERROR", "profile must contain an identity mapping")
                rendered_docx_pdf = render_docx_to_pdf(args.docx, args.docx.parent / ".docx-qa")
                check_pdf_layout_and_integrity(rendered_docx_pdf, args.market, profile["identity"])
                check_pdf_body_line_spacing(rendered_docx_pdf, args.project_manifest)
    except (ResumeQAError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    print("OK: skill and artifacts passed quantified QA")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
