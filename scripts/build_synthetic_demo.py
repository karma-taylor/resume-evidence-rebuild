#!/usr/bin/env python3
"""Create a synthetic, one-page resume pair for local smoke testing only."""
from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(29, 95, 167)
NAVY = RGBColor(16, 42, 67)
FONT = "Arial"


def add_run(paragraph, text, size, bold=False, color=NAVY):
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(.45); section.bottom_margin = Inches(.45)
    section.left_margin = Inches(.55); section.right_margin = Inches(.55)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(3)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(p, "Sample Candidate", 20, True)
    p = doc.add_paragraph(); add_run(p, "Enterprise AI Project Manager | RAG / Agent Skills / Harness", 11, True, BLUE)
    p = doc.add_paragraph(); add_run(p, "sample@example.com | Shanghai | Open to travel", 8.5, color=NAVY)
    for title, bullets in [
        ("Profile", ["Enterprise delivery professional who turns process discovery into scoped AI proofs of concept, evaluations, and handover-ready releases."]),
        ("Experience", [
            "Sample Company | Project Manager: led discovery, scope definition, cross-functional coordination, risk tracking, acceptance criteria, and delivery retrospectives.",
            "Managed the cadence from requirements clarification through prototype validation and release handover, with explicit data-access, human-review, and exception boundaries.",
        ]),
        ("Selected Projects", [
            "Enterprise Hybrid RAG: designed ACL-first filtering, BM25 plus dense retrieval, and RRF fusion; recorded Hit@5 on a synthetic fixed evaluation set and used regression checks to detect retrieval quality drift.",
            "Business Agent Harness: separated policy validation, structured output, failure handling, and audit logging into verifiable components while retaining an explicit human-review boundary.",
            "Resource Planning Tool: packaged people allocation, time-conflict prevention, and data recovery into a reusable project-planning workflow.",
        ]),
        ("Education & Certifications", ["Sample University, Master's degree | Generative AI Application Engineer (synthetic example)"]),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4)
        add_run(p, title, 11, True, BLUE)
        for bullet in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(.16)
            p.paragraph_format.first_line_indent = Inches(-.12)
            add_run(p, bullet, 8.8)
    doc.core_properties.author = "resume-evidence-rebuild synthetic fixture"
    doc.save(args.out)


if __name__ == "__main__":
    main()
