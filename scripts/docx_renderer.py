#!/usr/bin/env python3
"""Optional editable DOCX renderer; Typst remains the PDF layout authority."""
from __future__ import annotations

import argparse
import hashlib
import json
import sys
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from design_tokens import load_theme
from validate_resume_artifacts import (
    ResumeQAError,
    atomic_write_json,
    quarantine_artifacts,
    run_docx_delivery_gate,
)


DOCX_CJK_FONT = "Microsoft YaHei"
CONTENT_BOUNDS = {
    "normal": (40, 50),
    "compressed": (30, 40),
    "expanded": (50, 130),
}
BODY_LINE_SPACING = 1.4
BULLET_LINE_SPACING = 1.3
COMPACT_PARAGRAPH_AFTER_PT = 0.5
BULLET_PARAGRAPH_AFTER_PT = 5.0


def set_compact_paragraph(paragraph: Any, *, after_pt: float = COMPACT_PARAGRAPH_AFTER_PT) -> None:
    """Apply the fixed 1.4x body rhythm; reflow never changes it."""
    paragraph.paragraph_format.line_spacing = BODY_LINE_SPACING
    paragraph.paragraph_format.space_after = Pt(after_pt)
    paragraph.paragraph_format.space_before = Pt(0)


def set_bullet_paragraph(paragraph: Any, *, after_pt: float = BULLET_PARAGRAPH_AFTER_PT) -> None:
    """Apply 1.3x wrapped-line spacing plus a 5pt inter-bullet gap."""
    set_compact_paragraph(paragraph, after_pt=after_pt)
    paragraph.paragraph_format.line_spacing = BULLET_LINE_SPACING


def set_run_font(run: Any, size: float | None = None) -> None:
    """Use an explicit macOS CJK font so LibreOffice renders Chinese text."""
    run.font.name = DOCX_CJK_FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)


def color(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value.removeprefix("#"))


def set_run_color(run: Any, hex_value: str) -> None:
    run.font.color.rgb = color(hex_value)


def add_section_heading(doc: Document, label: str, accent: str, ink: str) -> None:
    paragraph = doc.add_paragraph()
    set_compact_paragraph(paragraph, after_pt=0)
    marker = paragraph.add_run("▌ ")
    marker.bold = True
    set_run_font(marker, 12)
    set_run_color(marker, accent)
    title = paragraph.add_run(label)
    title.bold = True
    set_run_font(title, 12)
    set_run_color(title, ink)


def add_title_date(
    doc: Document, title: str, date: str, *, ink: str, date_color: str,
    title_size: float = 11, before_pt: float = 0,
) -> Any:
    """Create a stable left-title/right-date row without tab drift."""
    paragraph = doc.add_paragraph()
    set_compact_paragraph(paragraph, after_pt=0)
    paragraph.paragraph_format.space_before = Pt(before_pt)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(17.6), WD_TAB_ALIGNMENT.RIGHT)
    title_run = paragraph.add_run(title)
    title_run.bold = True
    set_run_font(title_run, title_size)
    set_run_color(title_run, ink)
    paragraph.add_run("\t")
    date_run = paragraph.add_run(date)
    set_run_font(date_run, 10)
    set_run_color(date_run, "#000000")
    return paragraph


def add_project_rule(doc: Document, accent: str) -> Any:
    """Add a short blue rule beneath a project title."""
    paragraph = doc.add_paragraph()
    # Keep a visible title-to-body breathing space while leaving all textual
    # body paragraphs on the fixed 1.4x/0.5pt contract.
    set_compact_paragraph(paragraph, after_pt=3)
    paragraph.paragraph_format.right_indent = Cm(14.5)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), accent.removeprefix("#"))
    borders.append(bottom)
    p_pr.append(borders)
    return paragraph


def add_header_rule(doc: Document, accent: str) -> Any:
    """Add the full-width rule below the compact two-column identity row."""
    paragraph = doc.add_paragraph()
    set_compact_paragraph(paragraph, after_pt=0)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), accent.removeprefix("#"))
    borders.append(bottom)
    p_pr.append(borders)
    return paragraph


def remove_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def mark_header_table(table: Any) -> None:
    """Identify the allowed two-column identity row for structural QA."""
    description = OxmlElement("w:tblDescription")
    description.set(qn("w:val"), "resume-header")
    table._tbl.tblPr.append(description)


def load(path: Path) -> dict[str, Any]:
    with path.open(encoding="utf-8") as stream:
        value = yaml.safe_load(stream) if path.suffix in {".yaml", ".yml"} else json.load(stream)
    if not isinstance(value, dict):
        raise ValueError(f"{path} must contain an object")
    return value


def add_rich(paragraph: Any, text: str, phrases: list[str], size: float = 10) -> None:
    cursor = 0
    ordered = sorted(phrases, key=lambda item: text.find(item))
    for phrase in ordered:
        start = text.find(phrase, cursor)
        if start < 0:
            raise ValueError(f"declared bold phrase {phrase!r} is absent")
        if start > cursor:
            set_run_font(paragraph.add_run(text[cursor:start]), size)
        run = paragraph.add_run(phrase)
        run.bold = True
        set_run_font(run, size)
        cursor = start + len(phrase)
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size)


def business_display_text(detail: dict[str, Any], *, include_background: bool = True) -> str:
    structure = detail.get("business_structure") or {}
    if not all(isinstance(structure.get(key), dict) for key in ("business_difficulty", "solution_action", "quantified_result")):
        return str(detail["text"])
    stages = (("背景", "business_difficulty"), ("解决", "solution_action"), ("结果", "quantified_result"))
    if not include_background:
        stages = stages[1:]
    return "；".join(f"{label}：{structure[key].get('text', '')}" for label, key in stages)


def add_business_rich(paragraph: Any, detail: dict[str, Any], size: float = 10, *, include_background: bool = True) -> None:
    """Add explicit business-stage labels while preserving evidence text."""
    structure = detail.get("business_structure") or {}
    if not all(isinstance(structure.get(key), dict) for key in ("business_difficulty", "solution_action", "quantified_result")):
        add_rich(paragraph, str(detail["text"]), list(detail.get("bold_phrases_used", [])), size=size)
        return
    phrases = [str(item) for item in detail.get("bold_phrases_used", [])]
    stages = (("背景", "business_difficulty"), ("解决", "solution_action"), ("结果", "quantified_result"))
    if not include_background:
        stages = stages[1:]
    for index, (label, key) in enumerate(stages):
        if index:
            paragraph.add_run("；")
        label_run = paragraph.add_run(f"{label}：")
        label_run.bold = True
        set_run_font(label_run, size)
        add_rich(paragraph, str(structure[key].get("text", "")), [phrase for phrase in phrases if phrase in str(structure[key].get("text", ""))], size=size)


STAGE_LABELS = {"background": "背景", "solution": "解决", "result": "结果"}


def stage_display_text(detail: dict[str, Any]) -> str:
    return f"{STAGE_LABELS[str(detail.get('stage'))]}：{detail['text']}"


def add_stage_rich(paragraph: Any, detail: dict[str, Any], size: float = 10) -> None:
    label_run = paragraph.add_run(f"{STAGE_LABELS[str(detail.get('stage'))]}：")
    label_run.bold = True
    set_run_font(label_run, size)
    add_rich(paragraph, str(detail["text"]), list(detail.get("bold_phrases_used", [])), size=size)


def contains_cjk(text: str) -> bool:
    return any("\u3400" <= char <= "\u9fff" or "\uf900" <= char <= "\ufaff" for char in text)


def cjk_count(text: str) -> int:
    return sum("\u3400" <= char <= "\u9fff" or "\uf900" <= char <= "\ufaff" for char in text)


def _main_impl() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--profile", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--resume-plan", type=Path, required=True)
    parser.add_argument("--typeset-plan", type=Path, required=True)
    parser.add_argument("--theme-vars", "--design-tokens", dest="theme_vars", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--internal-delivery", action="store_true", help=argparse.SUPPRESS)
    args = parser.parse_args()
    profile, template, plan, typeset = map(load, (args.profile, args.template, args.resume_plan, args.typeset_plan))
    theme = load_theme(args.theme_vars)
    content_mode = str(typeset.get("content_mode", "normal"))
    if content_mode not in CONTENT_BOUNDS:
        raise ValueError(
            f"unsupported content_mode {content_mode!r}; expected normal, compressed, or expanded"
        )
    min_content_chars, max_content_chars = CONTENT_BOUNDS[content_mode]
    tokens = theme["tokens"]
    palette, hierarchy = tokens["palette"], tokens["hierarchy"]
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(1.7)
    normal = doc.styles["Normal"]
    normal.font.name, normal.font.size = DOCX_CJK_FONT, Pt(10)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_CJK_FONT)
    normal.font.color.rgb = color("#000000")
    normal.paragraph_format.line_spacing = BODY_LINE_SPACING
    normal.paragraph_format.space_after = Pt(COMPACT_PARAGRAPH_AFTER_PT)
    identity = profile["identity"]
    market = str(template.get("market", identity.get("market", ""))).upper()
    chinese = contains_cjk(template["target_role"])
    labels = ("技术能力", "工作经历", "项目经历", "教育与证书") if chinese else ("Technical Skills", "Work Experience", "Projects", "Education & Certifications")
    phone_label = "电话" if chinese else "Phone"
    email_label = "邮箱" if chinese else "Email"
    location_label = "地点" if chinese else "Location"
    portfolio_label = "作品集" if chinese else "Portfolio"
    contact_line_one = f"{phone_label}：{identity['phone']} | {email_label}：{identity['email']}"
    line_two_parts: list[str] = []
    if str(identity.get("location") or "").strip():
        line_two_parts.append(f"{location_label}：{identity['location']}")
    if str(identity.get("portfolio_url") or "").strip():
        line_two_parts.append(f"{portfolio_label}：{identity['portfolio_url']}")
    contact_lines = [contact_line_one]
    if line_two_parts:
        contact_lines.append(" | ".join(line_two_parts))
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_table.autofit = False
    remove_table_borders(header_table)
    mark_header_table(header_table)
    left_cell, right_cell = header_table.rows[0].cells
    left_cell.width = Cm(15.7)
    right_cell.width = Cm(2.30)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # Reuse the first paragraph in each cell so no empty spacer paragraph is
    # introduced before the identity block.
    title = left_cell.paragraphs[0]
    set_compact_paragraph(title, after_pt=0)
    title_run = title.add_run(identity["name"])
    title_run.bold = True; set_run_font(title_run, 20); set_run_color(title_run, palette["accent"])
    role = left_cell.add_paragraph(); set_compact_paragraph(role, after_pt=6)
    role_run = role.add_run(template["target_role"]); set_run_font(role_run, 11.5); set_run_color(role_run, palette["accent"])
    for line_index, line in enumerate(contact_lines):
        if not line:
            continue
        contact = left_cell.add_paragraph(); set_compact_paragraph(contact, after_pt=1 if line_index == 0 else 0)
        run = contact.add_run(line); set_run_font(run, 10); set_run_color(run, "#000000")
    if market == "CN":
        photo_path = str(identity.get("photo_path", "")).strip()
        photo = Path(photo_path).expanduser()
        if not photo_path or not photo.is_file():
            raise ValueError("CN resume requires an existing identity.photo_path")
        photo_paragraph = right_cell.paragraphs[0]
        photo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        photo_paragraph.paragraph_format.right_indent = Pt(8)
        photo_paragraph.add_run().add_picture(str(photo), width=Cm(2.30), height=Cm(3.07))
    add_header_rule(doc, palette["accent"])
    technical_skills = str(template.get("technical_skills") or template.get("summary", "")).strip()
    if technical_skills:
        add_section_heading(doc, labels[0], palette["accent"], palette["accent"])
        skills_para = doc.add_paragraph()
        set_compact_paragraph(skills_para)
        add_rich(skills_para, technical_skills, [str(item) for item in template.get("summary_bold_phrases", [])], size=10)
    add_section_heading(doc, labels[1], palette["accent"], palette["accent"])
    typeset_employment = {
        str(item.get("id")): item
        for item in typeset.get("employment", [])
        if isinstance(item, dict) and str(item.get("id", "")).strip()
    }
    expected_employment_ids = {f"employment-{index}" for index, _ in enumerate(profile.get("employment", []), 1)}
    if set(typeset_employment) != expected_employment_ids:
        raise ValueError("typeset plan must contain exactly one approved work-bullet set per employment entry")
    employment_manifest = []
    for index, item in enumerate(profile.get("employment", []), 1):
        employment_id = f"employment-{index}"
        add_title_date(
            doc, f"{item.get('employer', '')} | {item.get('title', '')}",
            f"{item.get('start', '')} – {item.get('end', '')}",
            ink=palette["accent"], date_color=hierarchy["date_color"], title_size=11,
        )
        details = typeset_employment[employment_id].get("bullets", [])
        work_min_chars, work_max_chars = CONTENT_BOUNDS["normal"]
        if (not isinstance(details, list) or not 4 <= len(details) <= 5
                or any(not isinstance(detail, dict) or not work_min_chars <= cjk_count(str(detail.get("text", ""))) <= work_max_chars for detail in details)):
            raise ValueError(f"each employment entry must contain 4-5 validated {work_min_chars}-{work_max_chars} CJK business bullets")
        for number, detail in enumerate(details, 1):
            para = doc.add_paragraph(style="List Bullet"); set_bullet_paragraph(para)
            # Work history uses direct business bullets; the project-only
            # 背景/解决/结果 labels are intentionally not rendered here.
            add_rich(para, str(detail["text"]), list(detail.get("bold_phrases_used", [])), size=10)
        employment_manifest.append({
            "id": employment_id,
            "name": f"{item.get('employer', '')} | {item.get('title', '')}",
            "bullets": [
                {"text": str(detail["text"]), "source_text": detail["text"], "bold_phrases": detail["bold_phrases_used"], "source_ingestion_ids": detail["source_ingestion_ids"], "derived_metric": (((detail.get("business_structure") or {}).get("quantified_result") or {}).get("derived_metric"))}
                for detail in details
            ],
        })
    projects = {item["id"]: item for item in plan["projects"]}
    add_section_heading(doc, labels[2], palette["accent"], palette["accent"])
    manifest = []
    for project_index, project in enumerate(typeset["projects"]):
        source = projects[project["id"]]
        stage_mode = bool(project["bullets"]) and all(isinstance(bullet, dict) and bullet.get("stage") for bullet in project["bullets"])
        add_title_date(
            doc, source["title"], f"{source['start']} – {source['end']}",
            ink=palette["accent"], date_color=hierarchy["date_color"], title_size=11,
            before_pt=6 if project_index else 0,
        )
        add_project_rule(doc, palette["accent"])
        overview_data = project.get("overview") or {}
        overview = str(overview_data.get("text") or "").strip()
        if overview and not stage_mode:
            para = doc.add_paragraph()
            set_bullet_paragraph(para)
            overview_run = para.add_run(f"背景：{overview}"); set_run_font(overview_run, 9); set_run_color(overview_run, "#000000")
        bullets = []
        for number, bullet in enumerate(project["bullets"], 1):
            para = doc.add_paragraph(style="List Bullet")
            set_bullet_paragraph(para)
            if stage_mode:
                add_stage_rich(para, bullet)
                display_text = stage_display_text(bullet)
            else:
                add_business_rich(para, bullet, include_background=False)
                display_text = business_display_text(bullet, include_background=False)
            derived_metric = bullet.get("derived_metric") if stage_mode else (((bullet.get("business_structure") or {}).get("quantified_result") or {}).get("derived_metric"))
            bullets.append({"element_id": f"project.{project['id']}.bullet.{number}", "text": display_text, "source_text": bullet["text"], "bold_phrases": bullet["bold_phrases_used"], "source_claim_ids": bullet["source_claim_ids"], "derived_metric": derived_metric})
        manifest.append({"id": project["id"], "name": source["title"], "overview": (None if stage_mode else (f"背景：{overview}" if overview else None)), "overview_source_claim_ids": overview_data.get("source_claim_ids", []) if not stage_mode else [], "bullets": bullets})
    add_section_heading(doc, labels[3], palette["accent"], palette["accent"])
    education = doc.add_paragraph(" · ".join(f"{entry.get('school', '')} {entry.get('degree', '')} {entry.get('start', '')}–{entry.get('end', '')}".strip() for entry in profile.get("education", [])))
    set_compact_paragraph(education)
    certs = doc.add_paragraph()
    set_compact_paragraph(certs)
    for index, cert in enumerate(profile.get("certifications", [])):
        if index: certs.add_run(" · ")
        run = certs.add_run(cert); run.bold = True
    # Paragraph styles are not consistently inherited by LibreOffice for CJK
    # glyphs.  Stamp the selected CJK font onto every generated body run.
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run)
    doc.core_properties.subject = f"resume-theme:{theme['variant_id']}"
    args.output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.output)
    manifest_path = args.output.parent / "docx-project-manifest.json"
    manifest_payload = {
        "projects": manifest,
        "employment": employment_manifest,
        "content_mode": content_mode,
        "theme_variant": theme["variant_id"],
        "provenance": {
            "renderer": "canonical_docx_renderer",
            "renderer_sha256": hashlib.sha256(Path(__file__).read_bytes()).hexdigest(),
            "profile_sha256": hashlib.sha256(args.profile.read_bytes()).hexdigest(),
            "template_sha256": hashlib.sha256(args.template.read_bytes()).hexdigest(),
            "resume_plan_sha256": hashlib.sha256(args.resume_plan.read_bytes()).hexdigest(),
            "typeset_plan_sha256": hashlib.sha256(args.typeset_plan.read_bytes()).hexdigest(),
            "theme_vars_sha256": hashlib.sha256(args.theme_vars.read_bytes()).hexdigest(),
        },
    }
    atomic_write_json(manifest_path, manifest_payload)
    qa_dir = args.output.parent / ".docx-qa"
    try:
        qa = run_docx_delivery_gate(
            docx_path=args.output, manifest_path=manifest_path,
            theme_path=args.theme_vars, profile_path=args.profile,
            market=market, qa_dir=qa_dir,
            renderer_path=Path(__file__),
        )
    except ResumeQAError as exc:
        if args.internal_delivery:
            raise
        quarantine = quarantine_artifacts(
            args.output.parent, (args.output, manifest_path),
            code=exc.code, detail=exc.detail, phase="docx_delivery",
        )
        print(json.dumps({"status": "delivery_gate_blocked", "error": str(exc), "quarantine": str(quarantine)}, ensure_ascii=False))
        return 3
    print(json.dumps({"status": "eligible_for_approval", "docx": str(args.output),
                      "project_manifest": str(manifest_path), "qa": qa}, ensure_ascii=False))
    return 0


def _cli_output_path(argv: list[str]) -> Path | None:
    try:
        index = argv.index("--output")
        return Path(argv[index + 1]).expanduser().resolve()
    except (ValueError, IndexError):
        return None


def main() -> int:
    """Ensure standalone DOCX failures enter the common SkillOpt boundary."""
    try:
        return _main_impl()
    except ResumeQAError:
        # Non-internal gate failures are already quarantined by _main_impl.
        raise
    except Exception as exc:
        if "--internal-delivery" not in sys.argv:
            output = _cli_output_path(sys.argv[1:])
            if output is not None:
                quarantine_artifacts(
                    output.parent, (output, output.parent / "docx-project-manifest.json"),
                    code="DOCX_RENDER_ERROR", detail=str(exc), phase="docx_delivery",
                )
        raise


if __name__ == "__main__":
    raise SystemExit(main())
