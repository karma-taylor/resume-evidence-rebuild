#!/usr/bin/env python3
"""Scan authorized resume material into an isolated inbox, then explicitly approve a merge."""
from __future__ import annotations

import argparse
import hashlib
import re
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Literal

import yaml
from docx import Document
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader


class SourceDocument(BaseModel):
    filename: str
    hash: str = Field(pattern=r"^[a-f0-9]{64}$")


class CandidateData(BaseModel):
    text: str = Field(min_length=1)
    inferred_type: Literal["context", "metric", "architecture", "control", "delivery"]
    flag: str | None = None


class PendingIngestion(BaseModel):
    ingestion_id: str
    status: Literal["pending", "approved", "rejected"]
    source_document: SourceDocument
    matched_employer: str
    locator: str
    candidate_data: list[CandidateData] = Field(min_length=1)


class Inbox(BaseModel):
    schema_version: str = "1.0"
    pending_ingestions: list[PendingIngestion] = Field(default_factory=list)


def read_inbox(path: Path) -> Inbox:
    if not path.exists():
        return Inbox()
    payload = yaml.safe_load(path.read_text(encoding="utf-8"))
    return Inbox.model_validate(payload or {})


def write_inbox(path: Path, inbox: Inbox) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(yaml.safe_dump(inbox.model_dump(mode="json"), allow_unicode=True, sort_keys=False), encoding="utf-8")


def source_hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def source_lines(path: Path) -> list[tuple[str, str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return [(line, f"Page {number}") for number, page in enumerate(reader.pages, 1) for line in (page.extract_text() or "").splitlines()]
    if suffix == ".docx":
        document = Document(str(path))
        return [(paragraph.text, f"Paragraph {number}") for number, paragraph in enumerate(document.paragraphs, 1)]
    if suffix in {".md", ".txt", ".yaml", ".yml"}:
        return [(line, f"Line {number}") for number, line in enumerate(path.read_text(encoding="utf-8").splitlines(), 1)]
    raise ValueError("supported source formats are PDF, DOCX, Markdown, text, and YAML")


def clean_line(line: str) -> str:
    return re.sub(r"^[\s•·*－—-]+", "", line).strip()


def infer_type(text: str) -> str:
    # A metric classification is intentionally narrow.  Words such as
    # “提升” without an actual number are delivery claims, not a license to
    # invent a quantified result at the end of a rendered bullet.
    if re.search(r"[0-9０-９]+(?:[.,．][0-9０-９]+)?\s*[%％]", text):
        return "metric"
    if re.search(r"架构|重构|检索|召回|Reranker|ACL|模块", text, re.IGNORECASE):
        return "architecture"
    if re.search(r"权限|校验|复核|审批|隔离|风控|治理|拦截", text):
        return "control"
    if re.search(r"问题|难点|风险|低效|遗漏|越权|瓶颈|需求", text):
        return "context"
    return "delivery"


def scan(args: argparse.Namespace) -> None:
    source = args.source.resolve()
    if not source.is_file():
        raise ValueError(f"source does not exist: {source}")
    inbox = read_inbox(args.inbox)
    if any(item.ingestion_id == args.ingestion_id for item in inbox.pending_ingestions):
        raise ValueError(f"ingestion_id already exists: {args.ingestion_id}")
    candidates = []
    for line, locator in source_lines(source):
        text = clean_line(line)
        if len(text) < 12 or text.endswith(("工作经历", "项目经历", "教育经历")):
            continue
        flag = "contains_independent_dev_claim" if any(token in text for token in ("独立", "0到1", "首个")) else None
        candidates.append((CandidateData(text=text, inferred_type=infer_type(text), flag=flag), locator))
    if not candidates:
        raise ValueError("no usable candidate statements found; upload a text-selectable document or enter facts manually")
    # Keep every candidate tied to a precise source locator; different page
    # locations become separate inbox entries rather than losing provenance.
    for index, (candidate, locator) in enumerate(candidates, 1):
        inbox.pending_ingestions.append(PendingIngestion(
            ingestion_id=f"{args.ingestion_id}_{index:02d}",
            status="pending",
            source_document=SourceDocument(filename=source.name, hash=source_hash(source)),
            matched_employer=args.employer,
            locator=locator,
            candidate_data=[candidate],
        ))
    write_inbox(args.inbox, inbox)
    print(f"OK: wrote {len(candidates)} pending candidates to {args.inbox}")


def approve(args: argparse.Namespace) -> None:
    inbox = read_inbox(args.inbox)
    selected = [item for item in inbox.pending_ingestions if item.ingestion_id in set(args.ingestion_ids)]
    if len(selected) != len(set(args.ingestion_ids)):
        raise ValueError("one or more ingestion IDs are absent")
    if any(item.status != "pending" for item in selected):
        raise ValueError("only pending ingestion entries can be approved")
    profile = yaml.safe_load(args.profile.read_text(encoding="utf-8"))
    if not isinstance(profile, dict) or not isinstance(profile.get("employment"), list):
        raise ValueError("profile must contain employment entries")
    employer = selected[0].matched_employer
    if any(item.matched_employer != employer for item in selected):
        raise ValueError("approve one employer per command")
    matching = [entry for entry in profile["employment"] if entry.get("employer") == employer]
    if len(matching) != 1:
        raise ValueError(f"exactly one profile employment entry must match {employer!r}")
    additions = [candidate.text for item in selected for candidate in item.candidate_data]
    raw_highlights = matching[0].get("highlights", [])
    if not isinstance(raw_highlights, list):
        raise ValueError("existing work highlights must be a list; migrate the legacy profile before approval")
    existing = [item for item in raw_highlights if isinstance(item, dict) and str(item.get("text", "")).strip()]
    if len(existing) != len(raw_highlights):
        raise ValueError("existing work highlights use the legacy bare-string format; confirm and migrate them before approval")
    if any(not isinstance(item.get("source_hash"), str) or len(str(item["source_hash"])) < 32 for item in existing):
        raise ValueError("existing work highlights use the legacy provenance format; confirm and migrate them before approval")
    existing_texts = {str(item["text"]).strip() for item in existing}
    approved_at = datetime.now(timezone.utc).isoformat()
    merged = existing + [
        {"text": candidate.text, "source_ingestion_id": item.ingestion_id, "approved_at": approved_at,
         "source_hash": item.source_document.hash}
        for item in selected for candidate in item.candidate_data
        if candidate.text not in existing_texts
    ]
    if not 4 <= len(merged) <= 5:
        raise ValueError("approval would not yield exactly 4-5 authorized work facts; select a valid set or provide more evidence")
    matching[0]["highlights"] = merged
    args.profile.write_text(yaml.safe_dump(profile, allow_unicode=True, sort_keys=False), encoding="utf-8")
    for item in selected:
        item.status = "approved"
    write_inbox(args.inbox, inbox)
    print(f"OK: approved {len(selected)} candidate(s) into {employer!r}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    subcommands = parser.add_subparsers(dest="command", required=True)
    scan_parser = subcommands.add_parser("scan")
    scan_parser.add_argument("--source", type=Path, required=True)
    scan_parser.add_argument("--employer", required=True)
    scan_parser.add_argument("--inbox", type=Path, required=True)
    scan_parser.add_argument("--ingestion-id", default=f"ing_{datetime.now(timezone.utc):%Y%m%d_%H%M%S}")
    scan_parser.set_defaults(handler=scan)
    approve_parser = subcommands.add_parser("approve")
    approve_parser.add_argument("--inbox", type=Path, required=True)
    approve_parser.add_argument("--profile", type=Path, required=True)
    approve_parser.add_argument("--ingestion-ids", nargs="+", required=True)
    approve_parser.set_defaults(handler=approve)
    args = parser.parse_args()
    try:
        args.handler(args)
    except (OSError, ValueError, ValidationError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
