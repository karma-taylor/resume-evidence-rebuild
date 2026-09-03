from __future__ import annotations

import sys
import json
import hashlib
import shutil
import subprocess
from pathlib import Path

import pytest
from pydantic import ValidationError
import yaml
from docx import Document
from docx.shared import Pt

sys.path.insert(0, str(Path(__file__).parents[1] / "scripts"))
from build_resume import (  # noqa: E402
    AssertionFragment, BusinessBulletStructure, BusinessSegment, Claim, DerivedMetric, EmploymentAssertion,
    EmploymentBusinessBulletStructure, EmploymentBusinessSegment, Employment, EvidenceGateError, NeedsUserInputError, Identity,
    Overview, Profile, Project, ProjectStageBullet, ResumePlan, ResumePlanProject, TypesetBullet, TypesetEmployment,
    TypesetEmploymentBullet, TypesetPlan, TypesetProject, WorkHighlight,
    RetryableContractError,
    Template, cjk_count, resolve_project_selection,
    validate_agent_a, validate_agent_b, validate_employment_provenance, validate_typeset_employment,
    validate_derived_metric, validate_derived_employment_metric, derived_metric_value,
    quarantine_build_failure, append_content_recovery_trace,
    generate_employment_typeset, validate_project_business_readability,
)
import typst_renderer  # noqa: E402
import docx_renderer  # noqa: E402
from design_tokens import theme_payload  # noqa: E402
from typst_renderer import FROZEN_LAYOUTS  # noqa: E402
from docx_renderer import BODY_LINE_SPACING, BULLET_LINE_SPACING, BULLET_PARAGRAPH_AFTER_PT, set_compact_paragraph  # noqa: E402
from validate_resume_artifacts import (  # noqa: E402
    ResumeQAError, begin_render_transaction, check_docx_ooxml_spacing, promote_render_transaction,
    check_delivery_manifest, check_typst_compact_body_spacing, paragraph_spacing_is_compact,
    quarantine_render_transaction, MIN_MARGIN_PT, MARGIN_TOLERANCE_PT, margin_below_minimum,
)


def test_typst_escape_preserves_literal_double_hyphen():
    assert typst_renderer.esc("匿名项目--") == "匿名项目\\-\\-"


def test_plan_only_failure_still_enters_skillopt_quarantine(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("SKILLOPT_AUTO_ENABLED", "0")
    quarantine = quarantine_build_failure(
        tmp_path / "output", None, None,
        code="EVIDENCE_GATE_BLOCKED", detail="missing authorized claim", phase="evidence_gate",
    )
    assert (quarantine / "failed-manifest.json").is_file()
    event = json.loads((quarantine / "skillopt-event.json").read_text(encoding="utf-8"))
    assert event["auto_skillopt"]["entered"] is True
    assert event["error_code"] == "EVIDENCE_GATE_BLOCKED"


def test_content_recovery_trace_appends_compress_then_prune(tmp_path: Path):
    append_content_recovery_trace(tmp_path, {
        "status": "content_compression_applied", "content_mode": "compressed",
    })
    append_content_recovery_trace(tmp_path, {
        "status": "project_prune_applied", "removed_project_id": "lowest-ranked",
    })
    trace = json.loads((tmp_path / "content-recovery-trace.json").read_text(encoding="utf-8"))
    assert [entry["status"] for entry in trace["attempts"]] == [
        "content_compression_applied", "project_prune_applied",
    ]
    assert trace["status"] == "project_prune_applied"


def claim(project: str, claim_id: str, text: str, kind: str) -> Claim:
    return Claim(id=claim_id, text=text, source="fixture", scope=project,
                 confidence="verified", allowed_for_resume=True, kind=kind)


def plan_and_copy(*, numeric: bool = False, unsupported_overview: bool = False, source_override: str | None = None):
    projects, copy_projects = [], []
    for index in range(3):
        project_id = f"p-{index}"
        difficulty = "岗位资料查询存在越权和遗漏风险"
        action = "配置岗位权限与复核流程"
        result = "各岗位核心资料查找准确率高达97.5%"
        context_claim = claim(project_id, f"c-context-{index}", difficulty, "context")
        action_claim = claim(project_id, f"c-action-{index}", action, "architecture")
        result_claim = claim(project_id, f"c-result-{index}", source_override or result,
                             "metric" if numeric or source_override is None else "delivery")
        project_claims = [context_claim, action_claim, result_claim]
        projects.append(ResumePlanProject(id=project_id, title=project_id, start="2024", end="2025", tags=[],
                                          claim_ids=[item.id for item in project_claims], claims=project_claims))
        bullet_text = difficulty + action + result
        copy_projects.append(TypesetProject(
            id=project_id,
            overview=Overview(text="虚构全球战略" if unsupported_overview and index == 0 else bullet_text,
                              source_claim_ids=[item.id for item in project_claims],
                              assertions=[AssertionFragment(text="虚构全球战略", source_claim_id=context_claim.id)] if unsupported_overview and index == 0 else [
                                  AssertionFragment(text=difficulty, source_claim_id=context_claim.id),
                                  AssertionFragment(text=action, source_claim_id=action_claim.id),
                                  AssertionFragment(text=result, source_claim_id=result_claim.id),
                              ]),
            bullets=[TypesetBullet(text=bullet_text, bold_phrases_used=[result], terminal_bold_phrase=result,
                                  source_claim_ids=[item.id for item in project_claims], assertions=[
                                      AssertionFragment(text=difficulty, source_claim_id=context_claim.id),
                                      AssertionFragment(text=action, source_claim_id=action_claim.id),
                                      AssertionFragment(text=result, source_claim_id=result_claim.id),
                                  ],
                                  business_structure=BusinessBulletStructure(
                                      business_difficulty=BusinessSegment(text=difficulty, source_claim_id=context_claim.id),
                                      solution_action=BusinessSegment(text=action, source_claim_id=action_claim.id),
                                      quantified_result=BusinessSegment(text=result, source_claim_id=result_claim.id),
                                  )) for _ in range(3)],
        ))
    return ResumePlan(target_role="x", projects=projects), TypesetPlan(projects=copy_projects, employment=[])


def test_unbound_overview_is_evidence_blocked():
    plan, copy = plan_and_copy(unsupported_overview=True)
    with pytest.raises(EvidenceGateError, match="verbatim Claim support"):
        validate_agent_b(copy, plan)


def test_numeric_terminal_requires_metric_claim():
    plan, copy = plan_and_copy(source_override="岗位资料查询存在越权和遗漏风险配置岗位权限与复核流程各岗位核心资料查找准确率高达97.5%")
    with pytest.raises(EvidenceGateError, match="metric"):
        validate_agent_b(copy, plan)


def test_agent_a_cannot_change_authorized_project_header():
    plan, _ = plan_and_copy()
    profile = Profile(
        identity=Identity(name="测试", phone="13800138000", email="a@example.com", portfolio_url="https://example.com", market="CN"),
        education=[], employment=[], certifications=[],
        projects=[Project(id=item.id, title=item.title, start=item.start, end=item.end, tags=item.tags, claims=item.claims) for item in plan.projects],
    )
    from build_resume import Template
    template = Template(id="fixture", target_role="x", market="CN", project_ids=[item.id for item in plan.projects], layout={})
    plan.projects[0].title = "伪造项目名称"
    with pytest.raises(EvidenceGateError, match="project header"):
        validate_agent_a(plan, profile, template)


def test_project_bullet_requires_an_authorized_numeric_result():
    difficulty = "岗位资料查询存在越权和遗漏风险"
    action = "配置岗位权限规则并设置人工复核流程"
    result = "权限审查流程已完成上线"
    bullet = TypesetBullet(
        text=difficulty + action + result,
        bold_phrases_used=[result],
        terminal_bold_phrase=result,
        source_claim_ids=["context", "action", "delivery"],
        assertions=[
            AssertionFragment(text=difficulty, source_claim_id="context"),
            AssertionFragment(text=action, source_claim_id="action"),
            AssertionFragment(text=result, source_claim_id="delivery"),
        ],
        business_structure=BusinessBulletStructure(
            business_difficulty=BusinessSegment(text=difficulty, source_claim_id="context"),
            solution_action=BusinessSegment(text=action, source_claim_id="action"),
            quantified_result=BusinessSegment(text=result, source_claim_id="delivery"),
        ),
    )
    claims = [
        claim("p-1", "context", difficulty, "context"),
        claim("p-1", "action", action, "architecture"),
        claim("p-1", "delivery", result, "delivery"),
    ]
    plan = ResumePlan(target_role="x", projects=[
        ResumePlanProject(id=f"p-{index}", title=f"p-{index}", start="2024", end="2025", tags=[],
                          claim_ids=[item.id for item in claims], claims=claims)
        for index in range(3)
    ])
    copy = TypesetPlan(projects=[
        TypesetProject(id=f"p-{index}", overview=Overview(text=difficulty, source_claim_ids=["context"], assertions=[AssertionFragment(text=difficulty, source_claim_id="context")]), bullets=[bullet, bullet, bullet])
        for index in range(3)
    ], employment=[])
    with pytest.raises(NeedsUserInputError, match="numeric or derived efficiency metric"):
        validate_agent_b(copy, plan)


def test_renderers_make_business_stage_labels_explicit():
    detail = {
        "text": "原始连续文案",
        "bold_phrases_used": ["准确率97.5%"],
        "business_structure": {
            "business_difficulty": {"text": "岗位资料查询存在越权风险"},
            "solution_action": {"text": "配置岗位权限与复核流程"},
            "quantified_result": {"text": "准确率97.5%"},
        },
    }
    display = typst_renderer.business_display_text(detail)
    assert display == "背景：岗位资料查询存在越权风险；解决：配置岗位权限与复核流程；结果：准确率97.5%"
    project_display = typst_renderer.business_display_text(detail, include_background=False)
    assert project_display == "解决：配置岗位权限与复核流程；结果：准确率97.5%"
    markup = typst_renderer.business_rich_text(detail)
    assert all(label in markup for label in ("背景：", "解决：", "结果："))
    assert 'weight: "bold"' in markup


def test_docx_ooxml_spacing_rejects_single_spacing(tmp_path):
    path = tmp_path / "single.docx"
    doc = Document()
    for text in ("姓名", "目标职位", "电话：123 | 邮箱：x@y.com", "地点：深圳 | 作品集：https://x.com"):
        header = doc.add_paragraph(text)
        set_compact_paragraph(header)
    paragraph = doc.add_paragraph("正文段落")
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(0.5)
    doc.save(path)
    with pytest.raises(ResumeQAError, match="PARAGRAPH_SPACING_ERROR"):
        check_docx_ooxml_spacing(path)


def test_docx_ooxml_spacing_accepts_exact_physical_contract(tmp_path):
    path = tmp_path / "compact.docx"
    doc = Document()
    for text in ("姓名", "目标职位", "电话：123 | 邮箱：x@y.com", "地点：深圳 | 作品集：https://x.com"):
        header = doc.add_paragraph(text)
        set_compact_paragraph(header)
    paragraph = doc.add_paragraph("正文段落")
    set_compact_paragraph(paragraph)
    doc.save(path)
    result = check_docx_ooxml_spacing(path)
    assert result["actual_line_values"] == [280]
    assert result["actual_after_values"] == [10]


def test_failed_render_transaction_is_quarantined(tmp_path):
    output = tmp_path / "output"
    run_id, staging = begin_render_transaction(output)
    failed_pdf = staging / "resume.pdf"
    failed_pdf.write_bytes(b"%PDF-failed")
    quarantine = quarantine_render_transaction(
        output, run_id, staging, code="PARAGRAPH_SPACING_ERROR",
        detail="single spacing", phase="docx_delivery",
    )
    assert (quarantine / "resume.pdf").is_file()
    failed_manifest = json.loads((quarantine / "failed-manifest.json").read_text(encoding="utf-8"))
    assert failed_manifest["status"] == "failed"
    event = json.loads((quarantine / "skillopt-event.json").read_text(encoding="utf-8"))
    assert event["error_code"] == "PARAGRAPH_SPACING_ERROR"
    assert event["auto_skillopt"]["entered"] is True
    assert event["auto_skillopt"]["eligible"] is True
    assert event["auto_skillopt"]["route"] == "public_rule_candidate"
    assert not (output / "delivery-manifest.json").exists()


def test_successful_promotion_consumes_staging_copy(tmp_path):
    output = tmp_path / "output"
    _, staging = begin_render_transaction(output)
    (staging / "resume.pdf").write_bytes(b"%PDF-approved")
    promoted = promote_render_transaction(staging, output, names=("resume.pdf",))
    assert promoted == [output / "resume.pdf"]
    assert (output / "resume.pdf").read_bytes() == b"%PDF-approved"
    assert not (staging / "resume.pdf").exists()
    staging.rmdir()


def test_project_stage_mode_requires_exactly_background_solution_result():
    projects, staged_projects = [], []
    for index in range(3):
        project_id = f"stage-{index}"
        background = "业务背景" * 10
        solution = "解决动作" * 10
        result = "结果" * 20 + "97.5%"
        claims = [
            claim(project_id, f"{project_id}-background", background, "context"),
            claim(project_id, f"{project_id}-solution", solution, "architecture"),
            claim(project_id, f"{project_id}-result", result, "metric"),
        ]
        projects.append(ResumePlanProject(id=project_id, title=project_id, start="2024", end="2025", tags=[], claim_ids=[item.id for item in claims], claims=claims))
        staged_projects.append(TypesetProject(id=project_id, overview=None, bullets=[
            ProjectStageBullet(stage="background", text=background, bold_phrases_used=["业务背景"], source_claim_ids=[claims[0].id], assertions=[AssertionFragment(text=background, source_claim_id=claims[0].id)]),
            ProjectStageBullet(stage="solution", text=solution, bold_phrases_used=["解决动作"], source_claim_ids=[claims[1].id], assertions=[AssertionFragment(text=solution, source_claim_id=claims[1].id)]),
            ProjectStageBullet(stage="result", text=result, bold_phrases_used=["97.5%"], terminal_bold_phrase="97.5%", source_claim_ids=[claims[2].id], assertions=[AssertionFragment(text=result, source_claim_id=claims[2].id)]),
        ]))
    plan = ResumePlan(target_role="x", projects=projects)
    validate_agent_b(TypesetPlan(projects=staged_projects, employment=[]), plan)
    with pytest.raises(ValidationError, match="exactly background"):
        TypesetProject(id="bad", overview=None, bullets=[
            staged_projects[0].bullets[0], staged_projects[0].bullets[1], staged_projects[0].bullets[1],
        ])


def test_project_content_mode_allows_bounded_density_recovery_only():
    projects, staged_projects = [], []
    for index in range(3):
        project_id = f"expanded-{index}"
        background = "业务背景" * 15  # 60 CJK
        solution = "解决动作" * 15      # 60 CJK
        result = "结果" * 30 + "97.5%"  # 60 CJK plus a supported metric
        claims = [
            claim(project_id, f"{project_id}-background", background, "context"),
            claim(project_id, f"{project_id}-solution", solution, "architecture"),
            claim(project_id, f"{project_id}-result", result, "metric"),
        ]
        projects.append(ResumePlanProject(id=project_id, title=project_id, start="2024", end="2025", tags=[], claim_ids=[item.id for item in claims], claims=claims))
        staged_projects.append(TypesetProject(id=project_id, overview=None, bullets=[
            ProjectStageBullet(stage="background", text=background, bold_phrases_used=["业务背景"], source_claim_ids=[claims[0].id], assertions=[AssertionFragment(text=background, source_claim_id=claims[0].id)]),
            ProjectStageBullet(stage="solution", text=solution, bold_phrases_used=["解决动作"], source_claim_ids=[claims[1].id], assertions=[AssertionFragment(text=solution, source_claim_id=claims[1].id)]),
            ProjectStageBullet(stage="result", text=result, bold_phrases_used=["97.5%"], terminal_bold_phrase="97.5%", source_claim_ids=[claims[2].id], assertions=[AssertionFragment(text=result, source_claim_id=claims[2].id)]),
        ]))
    plan = ResumePlan(target_role="x", projects=projects)
    validate_agent_b(TypesetPlan(projects=staged_projects, employment=[], content_mode="expanded"), plan, require_project_stages=True)
    with pytest.raises(RetryableContractError, match="normal budget is 40-50"):
        validate_agent_b(TypesetPlan(projects=staged_projects, employment=[], content_mode="normal"), plan, require_project_stages=True)


def test_render_admission_rejects_legacy_per_bullet_three_part_format():
    plan, copy = plan_and_copy()
    with pytest.raises(EvidenceGateError, match="three project-stage bullets"):
        validate_agent_b(copy, plan, require_project_stages=True)


def test_derived_metric_is_recomputed_from_one_metric_claim():
    source = claim("p-1", "m-1", "日报处理时间由2小时降至10分钟", "metric")
    metric = DerivedMetric(
        operation="percentage_reduction", source_claim_ids=["m-1"],
        before_text="2小时", after_text="10分钟",
        before_value=120, after_value=10, precision=1,
    )
    assert derived_metric_value(metric) == pytest.approx(91.6666667)
    validate_derived_metric(
        metric=metric, result_text="处理耗时下降约91.7%",
        claims={source.id: source}, declared_claim_ids=[source.id], label="p-1 bullet",
    )


def test_derived_metric_rejects_model_invented_result_or_cross_claim_inputs():
    source = claim("p-1", "m-1", "日报处理时间由2小时降至10分钟", "metric")
    metric = DerivedMetric(
        operation="percentage_reduction", source_claim_ids=["m-1"],
        before_text="2小时", after_text="10分钟",
        before_value=120, after_value=10, precision=1,
    )
    with pytest.raises(EvidenceGateError, match="computed value"):
        validate_derived_metric(
            metric=metric, result_text="处理耗时下降约95%",
            claims={source.id: source}, declared_claim_ids=[source.id], label="p-1 bullet",
        )
    forged_inputs = metric.model_copy(update={"before_value": 999})
    with pytest.raises(EvidenceGateError, match="before_value"):
        validate_derived_metric(
            metric=forged_inputs, result_text="处理耗时下降约99.0%",
            claims={source.id: source}, declared_claim_ids=[source.id], label="p-1 bullet",
        )
    before_claim = claim("p-1", "m-before", "日报处理时间由2小时开始", "metric")
    after_claim = claim("p-1", "m-after", "日报处理时间降至10分钟", "metric")
    cross_metric = metric.model_copy(update={"source_claim_ids": ["m-before", "m-after"]})
    with pytest.raises(EvidenceGateError, match="one metric Claim"):
        validate_derived_metric(
            metric=cross_metric, result_text="处理耗时下降约91.7%",
            claims={before_claim.id: before_claim, after_claim.id: after_claim},
            declared_claim_ids=[before_claim.id, after_claim.id], label="p-1 bullet",
        )


def test_derived_employment_metric_uses_approved_inbox_sources():
    metric = DerivedMetric(
        operation="times_improvement", source_claim_ids=["ing-1"],
        before_text="2小时", after_text="10分钟",
        before_value=120, after_value=10, precision=1,
    )
    def supports(source_id: str, text: str, kind: str | None = None) -> bool:
        return source_id == "ing-1" and kind == "metric" and text in "日报处理时间由2小时降至10分钟"
    validate_derived_employment_metric(
        metric=metric, result_text="处理效率提升12.0倍",
        declared_source_ids={"ing-1"}, source_supports=supports, label="employment-1 bullet",
    )


def test_jd_selection_requires_current_local_evidence_for_three_projects(tmp_path):
    source = tmp_path / "project-notes.md"
    source.write_text("RAG 权限治理\nAgent 定时调度\n回归评测\n", encoding="utf-8")
    digest = hashlib.sha256(source.read_bytes()).hexdigest()
    plan, _ = plan_and_copy()
    profile = Profile(
        identity=Identity(name="测试", phone="13800138000", email="a@example.com", portfolio_url="https://example.com", market="CN"),
        education=[], employment=[], certifications=[],
        projects=[Project(id=item.id, title=item.title, start=item.start, end=item.end, tags=item.tags, claims=item.claims) for item in plan.projects],
    )
    template = Template(id="fixture", target_role="x", market="CN", layout={})
    brief_path, map_path = tmp_path / "jd-brief.json", tmp_path / "jd-evidence-map.json"
    jd_hash = "a" * 64
    brief_path.write_text(json.dumps({"schema_version": "1.0", "target_role": "x", "jd_text_sha256": jd_hash, "requirements": [
        {"id": "req", "text": "匹配能力", "keywords": ["RAG", "Agent", "回归"], "priority": "required"}
    ]}), encoding="utf-8")
    map_path.write_text(json.dumps({"schema_version": "1.0", "jd_text_sha256": jd_hash, "matches": [
        {"project_id": f"p-{index}", "requirement_id": "req", "path": str(source), "line_start": index + 1, "line_end": index + 1, "source_sha256": digest, "excerpt": ["RAG 权限治理", "Agent 定时调度", "回归评测"][index], "matched_keywords": [["RAG"], ["Agent"], ["回归"]][index]}
        for index in range(3)
    ]}), encoding="utf-8")
    selection = resolve_project_selection(profile=profile, template=template, jd_brief_path=brief_path, jd_evidence_map_path=map_path)
    assert selection.mode == "jd"
    assert selection.project_ids == ["p-0", "p-1", "p-2"]
    source.write_text("内容已改变", encoding="utf-8")
    with pytest.raises(EvidenceGateError, match="source changed"):
        resolve_project_selection(profile=profile, template=template, jd_brief_path=brief_path, jd_evidence_map_path=map_path)


def test_layout_values_are_frozen_numeric_states():
    assert FROZEN_LAYOUTS["normal"]["module_gap"] == 7.0
    assert all(isinstance(value, float) and value >= 0 for layout in FROZEN_LAYOUTS.values() for value in layout.values())


def test_pdf_margin_boundary_tolerates_conversion_noise_but_rejects_real_violation():
    assert MIN_MARGIN_PT == pytest.approx(36.0)
    assert not margin_below_minimum(MIN_MARGIN_PT)
    assert not margin_below_minimum(MIN_MARGIN_PT - 1e-6)
    assert margin_below_minimum(MIN_MARGIN_PT - MARGIN_TOLERANCE_PT - 0.01)


def test_template_section_order_and_skills_alias_are_stable():
    template = Template(id="fixture", target_role="x", market="CN", project_ids=["p-0", "p-1", "p-2"], layout={})
    assert template.sections == ["profile", "technical-skills", "employment", "projects", "education-certifications"]
    assert Template(id="fixture", target_role="x", market="CN", project_ids=["p-0", "p-1", "p-2"], layout={}, technical_skills="RAG｜Agent").technical_skills == "RAG｜Agent"
    with pytest.raises(ValidationError, match="template sections"):
        Template(id="fixture", target_role="x", market="CN", project_ids=["p-0", "p-1", "p-2"], layout={}, sections=["profile", "employment", "projects", "education-certifications"])


def test_typst_artifact_cannot_silently_restore_sparse_paragraph_spacing(tmp_path: Path):
    source = tmp_path / "resume.typ"
    source.write_text('#let layout = json("layout_vars.json")\n#set text(top-edge: 0.8em, bottom-edge: -0.2em)\n#set par(leading: 0.1em, spacing: 8pt)\n', encoding="utf-8")
    with pytest.raises(ResumeQAError, match="PARAGRAPH_SPACING_ERROR"):
        check_typst_compact_body_spacing(source)
    source.write_text(
        '#let layout = json("layout_vars.json")\n'
        '#let layout-len(key) = layout.at("spacing").at(key) * 1pt\n'
        '#set text(top-edge: 0.8em, bottom-edge: -0.2em)\n'
        '#set par(leading: 0.4em, spacing: 0.5pt)\n'
        '#let bullet(body) = { set par(leading: 0.3em, spacing: 5pt); body }\n'
        '#v(layout-len("header_to_first_module"))\n#v(layout-len("module_gap"))\n'
        '#v(layout-len("project_gap"))\n#v(layout-len("title_to_overview"))\n#v(layout-len("overview_to_bullet"))\n',
        encoding="utf-8",
    )
    check_typst_compact_body_spacing(source)
    source.write_text(source.read_text(encoding="utf-8") + '#set par(leading: 20pt, spacing: 12pt)\n', encoding="utf-8")
    with pytest.raises(ResumeQAError, match="PARAGRAPH_SPACING_ERROR"):
        check_typst_compact_body_spacing(source)
    source.write_text(source.read_text(encoding="utf-8").replace('#set par(leading: 20pt, spacing: 12pt)\n', '') + '#v(12pt)\n', encoding="utf-8")
    with pytest.raises(ResumeQAError, match="PARAGRAPH_SPACING_ERROR"):
        check_typst_compact_body_spacing(source)


def test_docx_body_uses_fixed_one_point_four_line_spacing(tmp_path: Path):
    document = Document()
    paragraph = document.add_paragraph("测试正文")
    set_compact_paragraph(paragraph)
    path = tmp_path / "spacing.docx"
    document.save(path)
    rendered = Document(path).paragraphs[0]
    assert BODY_LINE_SPACING == 1.4
    assert BULLET_LINE_SPACING == 1.3
    assert BULLET_PARAGRAPH_AFTER_PT == 2.0
    assert paragraph_spacing_is_compact(rendered)


def test_typst_one_point_four_line_height_directive_compiles(tmp_path: Path):
    binary = shutil.which("typst") or ""
    if not Path(binary).is_file():
        pytest.skip("Typst binary is unavailable")
    source = tmp_path / "line-height.typ"
    pdf = tmp_path / "line-height.pdf"
    source.write_text(
        '#set text(font: "Microsoft YaHei", size: 10pt)\n'
        '#set text(top-edge: 0.8em, bottom-edge: -0.2em)\n'
        '#set par(leading: 0.4em, spacing: 0.5pt)\n'
        '第一行正文用于验证。\\\n第二行正文用于验证。\n',
        encoding="utf-8",
    )
    subprocess.run([binary, "compile", str(source), str(pdf)], check=True, capture_output=True)
    assert pdf.is_file() and pdf.stat().st_size > 0


def test_business_segments_cannot_reuse_terminal_metric_for_every_role():
    difficulty = "岗位资料查询存在越权和遗漏风险"
    action = "配置岗位权限与复核流程"
    result = "各岗位核心资料查找准确率高达97.5%"
    with pytest.raises(ValidationError, match="distinct"):
        TypesetBullet(
            text=difficulty + action + result,
            bold_phrases_used=[result], terminal_bold_phrase=result,
            source_claim_ids=["c-1"], assertions=[AssertionFragment(text=difficulty + action + result, source_claim_id="c-1")],
            business_structure=BusinessBulletStructure(
                business_difficulty=BusinessSegment(text=result, source_claim_id="c-1"),
                solution_action=BusinessSegment(text=result, source_claim_id="c-1"),
                quantified_result=BusinessSegment(text=result, source_claim_id="c-1"),
            ),
        )


def test_business_segment_claim_must_belong_to_the_current_project():
    plan, copy = plan_and_copy()
    first = copy.projects[0].bullets[0]
    first.business_structure.solution_action.source_claim_id = "c-1"
    with pytest.raises(EvidenceGateError, match="cross-project"):
        validate_agent_b(copy, plan)


def employment_fixture(tmp_path: Path):
    difficulty = "岗位资料查询存在越权和遗漏风险"
    action = "配置岗位权限与复核流程"
    result = "各岗位核心资料查找准确率高达97.5%"
    text = difficulty + action + result
    assert cjk_count(text) == 40
    source_ids = [f"ing-{number}" for number in range(1, 5)]
    profile = Profile(
        identity=Identity(name="测试", phone="13800138000", email="a@example.com", portfolio_url="https://example.com", market="CN"),
        education=[], certifications=[], projects=[],
        employment=[Employment(employer="测试公司", title="项目经理", start="2024", end="2025", highlights=[
            WorkHighlight(text=text, source_ingestion_id=source_id, approved_at="2026-08-30T00:00:00Z", source_hash="a" * 64)
            for source_id in source_ids
        ])],
    )
    inbox = tmp_path / "ingestion_inbox.yaml"
    inbox.write_text(yaml.safe_dump({"schema_version": "1.0", "pending_ingestions": [
        {"ingestion_id": source_id, "status": "approved", "source_document": {"filename": "resume.docx", "hash": "a" * 64}, "matched_employer": "测试公司", "locator": "Paragraph 1", "candidate_data": [
            {"text": text, "inferred_type": "delivery"},
            {"text": difficulty, "inferred_type": "context"},
            {"text": action, "inferred_type": "architecture"},
            {"text": result, "inferred_type": "metric"},
        ]}
        for source_id in source_ids
    ]}, allow_unicode=True), encoding="utf-8")
    separators = ("", "，", "；", "、")
    bullet = lambda source_id, separator: TypesetEmploymentBullet(
        text=difficulty + separator + action + result, bold_phrases_used=[result], terminal_bold_phrase=result,
        source_ingestion_ids=[source_id],
        assertions=[
            EmploymentAssertion(text=difficulty, source_ingestion_id=source_id),
            EmploymentAssertion(text=action, source_ingestion_id=source_id),
            EmploymentAssertion(text=result, source_ingestion_id=source_id),
        ],
        business_structure=EmploymentBusinessBulletStructure(
            business_difficulty=EmploymentBusinessSegment(text=difficulty, source_ingestion_id=source_id),
            solution_action=EmploymentBusinessSegment(text=action, source_ingestion_id=source_id),
            quantified_result=EmploymentBusinessSegment(text=result, source_ingestion_id=source_id),
        ),
    )
    return profile, inbox, [bullet(source_id, separator) for source_id, separator in zip(source_ids, separators)]


def test_employment_bullets_may_recompose_only_approved_source_facts(tmp_path: Path):
    profile, inbox, bullets = employment_fixture(tmp_path)
    _, copy = plan_and_copy()
    copy.employment = [TypesetEmployment(id="employment-1", bullets=bullets)]
    validate_typeset_employment(copy, profile, inbox)


def test_work_bullet_can_be_direct_business_sentence_without_project_stage_structure(tmp_path: Path):
    profile, inbox, bullets = employment_fixture(tmp_path)
    source = bullets[0]
    direct = TypesetEmploymentBullet(
        text=source.text,
        bold_phrases_used=source.bold_phrases_used,
        terminal_bold_phrase=source.terminal_bold_phrase,
        source_ingestion_ids=source.source_ingestion_ids,
        assertions=source.assertions,
    )
    assert direct.business_structure is None
    _, copy = plan_and_copy()
    copy.employment = [TypesetEmployment(id="employment-1", bullets=[direct, *bullets[1:]])]
    validate_typeset_employment(copy, profile, inbox)


def test_work_bullet_rejects_technical_terms_outside_action_fragment(tmp_path: Path):
    profile, inbox, bullets = employment_fixture(tmp_path)
    payload = yaml.safe_load(inbox.read_text(encoding="utf-8"))
    source = payload["pending_ingestions"][0]
    context = "面向多部门岗位资料查询存在越权和遗漏风险"
    source["candidate_data"][1]["text"] = f"Python{context}"
    inbox.write_text(yaml.safe_dump(payload, allow_unicode=True), encoding="utf-8")
    result = "各岗位核心资料查找准确率高达97.5%"
    action = "配置岗位权限与复核流程"
    direct = TypesetEmploymentBullet(
        text=f"Python{context}；{action}{result}",
        bold_phrases_used=[result], terminal_bold_phrase=result,
        source_ingestion_ids=["ing-1"],
        assertions=[
            EmploymentAssertion(text=f"Python{context}", source_ingestion_id="ing-1"),
            EmploymentAssertion(text=action, source_ingestion_id="ing-1"),
            EmploymentAssertion(text=result, source_ingestion_id="ing-1"),
        ],
    )
    _, copy = plan_and_copy()
    copy.employment = [TypesetEmployment(id="employment-1", bullets=[direct, *bullets[1:]])]
    with pytest.raises(EvidenceGateError, match="TECHNICAL_TERM_PLACEMENT_ERROR"):
        validate_typeset_employment(copy, profile, inbox)


def test_automatic_work_generation_requires_business_context(tmp_path: Path):
    profile, inbox, _ = employment_fixture(tmp_path)
    payload = yaml.safe_load(inbox.read_text(encoding="utf-8"))
    for entry in payload["pending_ingestions"]:
        entry["candidate_data"] = [candidate for candidate in entry["candidate_data"] if candidate.get("inferred_type") != "context"]
    inbox.write_text(yaml.safe_dump(payload, allow_unicode=True), encoding="utf-8")
    with pytest.raises(EvidenceGateError, match="BUSINESS_CONTEXT_MISSING"):
        generate_employment_typeset(profile, inbox)


def test_project_background_rejects_technical_framing():
    with pytest.raises(EvidenceGateError, match="TECHNICAL_TERM_PLACEMENT_ERROR"):
        validate_project_business_readability(
            "RAG业务背景资料存在岗位越权和遗漏风险，需要按岗位区分访问权限并保留复核记录",
            "background",
            "p-1 background bullet",
        )


def test_project_solution_rejects_technical_term_stacking():
    with pytest.raises(EvidenceGateError, match="TECHNICAL_TERM_OVERLOAD"):
        validate_project_business_readability(
            "使用 Python FastAPI RAG ACL JWT 完成资料查询与权限控制",
            "solution",
            "p-1 solution bullet",
        )


def test_employment_bullet_rejects_foreign_source_fact(tmp_path: Path):
    profile, inbox, bullets = employment_fixture(tmp_path)
    bullets[0].source_ingestion_ids = ["not-approved"]
    _, copy = plan_and_copy()
    copy.employment = [TypesetEmployment(id="employment-1", bullets=bullets)]
    with pytest.raises(EvidenceGateError, match="unapproved or foreign"):
        validate_typeset_employment(copy, profile, inbox)


def test_work_highlight_source_hash_tampering_is_evidence_blocked(tmp_path: Path):
    profile, inbox, _ = employment_fixture(tmp_path)
    profile.employment[0].highlights[0].source_hash = "b" * 64
    with pytest.raises(EvidenceGateError, match="source hash"):
        validate_employment_provenance(profile, inbox)


def test_direct_renderer_requires_gate_before_writing_delivery_manifest(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    profile = tmp_path / "profile.yaml"
    template = tmp_path / "template.yaml"
    resume_plan = tmp_path / "resume-plan.json"
    typeset_plan = tmp_path / "typeset-plan.json"
    theme = tmp_path / "theme.json"
    layout = tmp_path / "layout.json"
    output = tmp_path / "output"
    profile.write_text(yaml.safe_dump({"identity": {"name": "测试", "phone": "13800138000", "email": "a@example.com", "portfolio_url": "https://example.com", "market": "FOREIGN"}, "employment": [], "education": [], "certifications": []}, allow_unicode=True), encoding="utf-8")
    template.write_text(yaml.safe_dump({"target_role": "测试岗位", "market": "FOREIGN", "layout": {"body_line_height_multiplier": 1.4}}, allow_unicode=True), encoding="utf-8")
    projects = [{"id": f"p{i}", "title": f"项目{i}", "start": "2024", "end": "2025"} for i in range(3)]
    resume_plan.write_text(json.dumps({"projects": projects}, ensure_ascii=False), encoding="utf-8")
    typeset_plan.write_text(json.dumps({"projects": [{"id": f"p{i}", "overview": {}, "bullets": [{"text": "测试内容", "bold_phrases_used": [], "source_claim_ids": []} for _ in range(3)]} for i in range(3)]}, ensure_ascii=False), encoding="utf-8")
    theme.write_text(json.dumps(theme_payload("executive_editorial_a")), encoding="utf-8")
    layout.write_text(json.dumps({"layout_state": "normal", "spacing": FROZEN_LAYOUTS["normal"]}), encoding="utf-8")

    def fake_compile(command, check):
        Path(command[-1]).write_bytes(b"%PDF-1.4\n")

    calls = []
    def pass_gate(**kwargs):
        calls.append(kwargs)
        return {"passed": True, "findings": []}

    class FakeValidated:
        def __init__(self, path: Path):
            self.path = path

        def model_dump(self, *, mode: str):
            return typst_renderer.load(self.path)

    def admit(**kwargs):
        admission_calls.append(kwargs)
        return tuple(FakeValidated(kwargs[key]) for key in (
            "profile_path", "template_path", "resume_plan_path", "typeset_plan_path",
        ))

    monkeypatch.setattr(typst_renderer.subprocess, "run", fake_compile)
    monkeypatch.setattr(typst_renderer.shutil, "which", lambda name: "/usr/bin/typst" if name == "typst" else None)
    monkeypatch.setattr(typst_renderer, "run_pdf_delivery_gate", pass_gate)
    admission_calls = []
    monkeypatch.setattr(typst_renderer, "load_and_validate_render_inputs", admit)
    monkeypatch.setattr(sys, "argv", ["typst_renderer.py", "--profile", str(profile), "--template", str(template), "--resume-plan", str(resume_plan), "--typeset-plan", str(typeset_plan), "--output-dir", str(output), "--layout-vars", str(layout), "--theme-vars", str(theme)])
    assert typst_renderer.main() == 0
    assert admission_calls and admission_calls[0]["profile_path"] == profile
    assert calls and calls[0]["allow_density"] is False
    assert (output / "delivery-manifest.json").is_file()

    def fail_gate(**_kwargs):
        raise ResumeQAError("PARAGRAPH_SPACING_ERROR", "sparse source")

    monkeypatch.setattr(typst_renderer, "run_pdf_delivery_gate", fail_gate)
    (output / "delivery-manifest.json").unlink()
    monkeypatch.setattr(sys, "argv", ["typst_renderer.py", "--profile", str(profile), "--template", str(template), "--resume-plan", str(resume_plan), "--typeset-plan", str(typeset_plan), "--output-dir", str(output), "--layout-vars", str(layout), "--theme-vars", str(theme)])
    with pytest.raises(ResumeQAError, match="PARAGRAPH_SPACING_ERROR"):
        typst_renderer.main()
    assert not (output / "resume.pdf").exists()
    assert not (output / "delivery-manifest.json").exists()


def test_docx_delivery_manifest_detects_post_qa_project_manifest_tampering(tmp_path: Path):
    files = {name: tmp_path / name for name in ("resume.docx", "project-manifest.json", "docx-project-manifest.json", "theme_vars.json")}
    for path in files.values():
        path.write_bytes(b"fixture")
    manifest = tmp_path / "docx-delivery-manifest.json"
    manifest.write_text(json.dumps({
        "status": "eligible_for_approval",
        "sha256": {
            "docx": hashlib.sha256(files["resume.docx"].read_bytes()).hexdigest(),
            "project_manifest": hashlib.sha256(files["project-manifest.json"].read_bytes()).hexdigest(),
            "docx_project_manifest": hashlib.sha256(files["docx-project-manifest.json"].read_bytes()).hexdigest(),
            "theme_vars": hashlib.sha256(files["theme_vars.json"].read_bytes()).hexdigest(),
        },
    }), encoding="utf-8")
    expected = {"docx": files["resume.docx"], "project_manifest": files["project-manifest.json"],
                "docx_project_manifest": files["docx-project-manifest.json"], "theme_vars": files["theme_vars.json"]}
    check_delivery_manifest(manifest, expected_paths=expected)
    files["docx-project-manifest.json"].write_bytes(b"tampered")
    with pytest.raises(ResumeQAError, match="docx_project_manifest"):
        check_delivery_manifest(manifest, expected_paths=expected)


def test_unbound_ascii_metric_injection_is_evidence_blocked():
    plan, copy = plan_and_copy()
    bullet = copy.projects[0].bullets[0]
    bullet.text = "999%" + bullet.text
    with pytest.raises(EvidenceGateError, match="unbound assertion text"):
        validate_agent_b(copy, plan)


def test_direct_renderer_admission_failure_precedes_typst_compile(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    from build_resume import EvidenceGateError
    compile_calls = []
    monkeypatch.setattr(typst_renderer, "load_and_validate_render_inputs",
                        lambda **_: (_ for _ in ()).throw(EvidenceGateError("EVIDENCE_GATE_BLOCKED: fixture")))
    monkeypatch.setattr(typst_renderer.subprocess, "run", lambda *_args, **_kwargs: compile_calls.append(True))
    paths = {name: tmp_path / name for name in ("profile.yaml", "template.yaml", "resume.json", "typeset.json", "theme.json", "layout.json")}
    for path in paths.values():
        path.write_text("{}", encoding="utf-8")
    monkeypatch.setattr(sys, "argv", [
        "typst_renderer.py", "--profile", str(paths["profile.yaml"]), "--template", str(paths["template.yaml"]),
        "--resume-plan", str(paths["resume.json"]), "--typeset-plan", str(paths["typeset.json"]),
        "--output-dir", str(tmp_path / "output"), "--layout-vars", str(paths["layout.json"]), "--theme-vars", str(paths["theme.json"]),
    ])
    with pytest.raises(EvidenceGateError, match="EVIDENCE_GATE_BLOCKED"):
        typst_renderer.main()
    assert not compile_calls


def test_standalone_renderer_preflight_errors_are_quarantined(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    def fail_typst() -> int:
        raise ValueError("synthetic typst preflight error")

    def fail_docx() -> int:
        raise ValueError("synthetic docx preflight error")

    typst_quarantine: list[dict] = []
    docx_quarantine: list[dict] = []
    monkeypatch.setattr(typst_renderer, "_main_impl", fail_typst)
    monkeypatch.setattr(docx_renderer, "_main_impl", fail_docx)
    monkeypatch.setattr(typst_renderer, "quarantine_artifacts", lambda *args, **kwargs: typst_quarantine.append(kwargs))
    monkeypatch.setattr(docx_renderer, "quarantine_artifacts", lambda *args, **kwargs: docx_quarantine.append(kwargs))
    monkeypatch.setattr(sys, "argv", ["typst_renderer.py", "--output-dir", str(tmp_path / "pdf")])
    with pytest.raises(ValueError, match="synthetic typst"):
        typst_renderer.main()
    monkeypatch.setattr(sys, "argv", ["docx_renderer.py", "--output", str(tmp_path / "docx" / "resume.docx")])
    with pytest.raises(ValueError, match="synthetic docx"):
        docx_renderer.main()
    assert typst_quarantine[0]["code"] == "TYPST_RENDER_ERROR"
    assert docx_quarantine[0]["code"] == "DOCX_RENDER_ERROR"
